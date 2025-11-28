"""
Generate Word document for MentorAid project overview
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Title
title = doc.add_heading("MentorAid — AI-based Student Dropout Prediction", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph(
    "Intelligent Early Warning System for Educational Institutions"
)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.italic = True

doc.add_paragraph()  # Spacing

# ========== THE PROBLEM ==========
doc.add_heading("The Problem: India's Dropout Crisis", 1)

crisis_intro = doc.add_paragraph(
    "India's school system supports nearly 25 crore students, but dropout remains a persistent national crisis. "
    "Over 5.4 million students drop out every year, with secondary and higher-secondary stages hit hardest."
)
crisis_intro.paragraph_format.space_after = Pt(8)

rajasthan_stats = [
    "Rajasthan Crisis: Nearly 9 lakh (900,000) children left school in 2023-24 — over 2,400 students per day",
    "Elementary Dropouts: 899,240 children (Grades 1–8) dropped out between 2022-2024 academic years",
    "Girls at Risk: Dropout rate exceeds 20% at age 16, higher than national average at every stage",
    "Higher Education: 20-25% dropout in undergraduate programs, higher in rural and low-income groups",
]
for stat in rajasthan_stats:
    doc.add_paragraph(stat, style="List Bullet")

doc.add_heading("Root Causes", 2)
causes = [
    "Poor attendance and academic performance",
    "Financial hardships, unpaid fees, and early marriage",
    "Fragmented data (attendance, marks, fees in separate spreadsheets)",
    "Lack of timely interventions or counseling",
    "Limited parent awareness and no early warning systems",
]
for cause in causes:
    doc.add_paragraph(cause, style="List Bullet")

doc.add_heading("Why It Matters", 2)
impact = doc.add_paragraph(
    "Each dropout affects family livelihoods, workforce strength, and perpetuates cycles of poverty and early marriage. "
    "For every year of secondary education missed, lifetime income potential drops by ₹60,000-₹1,20,000 per person. "
    "Traditional approaches focus on enrollment but fail to monitor and act on at-risk students due to fragmented data."
)
impact.paragraph_format.space_after = Pt(12)

# ========== PROJECT SUMMARY ==========
doc.add_heading("Our Solution: MentorAid", 1)
summary = doc.add_paragraph(
    "MentorAid is an AI-powered dropout prevention platform that predicts, prevents, and reduces student dropout at scale. "
    "It combines attendance, marks, fee data, and institutional records into one intelligent system. Upload CSV files or use "
    "individual risk predictor tools to get instant risk scores, confidence levels, and AI-generated intervention plans. "
    "With Google Sign-In, visual dashboards, chatbot assistant, and explainable AI insights, MentorAid makes early intervention "
    "practical and fast."
)
summary.paragraph_format.space_after = Pt(12)

key_point = doc.add_paragraph(
    "Use this to identify at-risk students, prioritize outreach, and measure intervention impact.",
    style="Intense Quote",
)

# ========== REAL-TIME APPLICATIONS ==========
doc.add_heading("Real-time Applications", 1)
applications = [
    (
        "Early Warning System",
        "Spot high-risk students from attendance and grade patterns",
    ),
    (
        "Bulk Batch Processing",
        "Upload class or school CSVs to predict risk for all students instantly",
    ),
    (
        "Teacher Dashboard",
        "Filter by class/department and take targeted action for students who need help",
    ),
    (
        "Counseling & Interventions",
        "Provides suggested remediation and records intervention history",
    ),
    (
        "Admin Reporting",
        "Aggregated stats and charts to monitor trends and program effectiveness",
    ),
    (
        "Chatbot Assistant",
        "Quick navigation, upload triggers, and short how-to guidance",
    ),
]

for app_name, app_desc in applications:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{app_name}: ").bold = True
    p.add_run(app_desc)

# ========== HOW IT WORKS ==========
doc.add_heading("How MentorAid Works (Step-by-Step)", 1)
steps = [
    (
        "Frontend Login",
        "Teacher/Admin logs in using Google Sign-In (OAuth) on the web app",
    ),
    (
        "Upload Data",
        "User uploads a CSV with student fields (attendance, marks, class, Nationality, fees, etc.) or uses the quick predictor form",
    ),
    (
        "API Call",
        "Frontend sends the file to the backend /api/predict/batch endpoint with the user token",
    ),
    (
        "Preprocessing",
        "Backend loads CSV with pandas, computes derived fields (attendance%, averageMarks), handles legacy issues (Nationality → Nacionality), normalizes categorical columns",
    ),
    (
        "Prediction",
        "Preprocessed features fed into Random Forest model (scikit-learn) returning prediction, confidence, and probabilities",
    ),
    (
        "Storage",
        "Results stored in database (MongoDB preferred; SQLite fallback). Probabilities serialized to JSON for SQLite compatibility",
    ),
    (
        "Response & UI",
        "API returns enriched rows. Frontend displays table, charts, student detail view with recommended interventions",
    ),
    (
        "Chatbot & Quick Actions",
        "In-app chatbot navigates to dashboard, triggers upload modal, or opens quick risk predictor. Uses Gemini API if configured, otherwise rule-based fallback",
    ),
]

for i, (step_name, step_desc) in enumerate(steps, 1):
    p = doc.add_paragraph(style="List Number")
    p.add_run(f"{step_name}: ").bold = True
    p.add_run(step_desc)

# ========== TECHNOLOGY STACK ==========
doc.add_heading("Technology Stack (Complete)", 1)

# Frontend
doc.add_heading("Frontend", 2)
frontend_items = [
    "React + Vite + TypeScript",
    "Styling: Tailwind CSS",
    "Components: Custom dashboard, file upload (drag/drop), charts",
    "Hosting: Netlify",
    "Environment Variables:",
    "  • VITE_API_URL — Backend base URL",
    "  • VITE_GOOGLE_CLIENT_ID — Google OAuth client ID",
    "  • VITE_GEMINI_API_KEY — Google Gemini API key (optional for chatbot)",
]
for item in frontend_items:
    doc.add_paragraph(item, style="List Bullet")

# Backend
doc.add_heading("Backend", 2)
backend_items = [
    "Python 3.11+ with Flask",
    "WSGI: Gunicorn (production on Render); Waitress (optional for local dev)",
    "Key Libraries:",
    "  • pandas, numpy — Data handling",
    "  • scikit-learn — Random Forest model",
    "  • imbalanced-learn — SMOTE / balancing utilities",
    "  • joblib — Saving/loading model artifacts",
    "  • flask-jwt-extended — JWT authentication",
    "  • flask-cors — Cross Origin Resource Sharing",
    "  • google-auth / google-auth-oauthlib — Google OAuth verification",
    "  • pymongo or SQLite — Database (MongoDB preferred; SQLite fallback)",
    "Hosting: Render (web service) or any WSGI host",
]
for item in backend_items:
    doc.add_paragraph(item, style="List Bullet")

# Database
doc.add_heading("Database", 2)
db_items = [
    "Primary: MongoDB (recommended for production)",
    "Fallback: SQLite (mentoraid.db) with migration logic to add missing columns",
    "Stored fields: student attributes, averageMarks, attendance, feeStatus, prediction, confidence, probabilities (JSON string in SQLite)",
]
for item in db_items:
    doc.add_paragraph(item, style="List Bullet")

# ML Model
doc.add_heading("Machine Learning Model & Tools", 2)
ml_items = [
    "Model: Random Forest classifier (scikit-learn)",
    "Training: Uses ~47 features (original + engineered). Important features include marks, attendance, fee/payment flags, engagement metrics",
    "Class balancing: SMOTE used during training to handle imbalance",
    "Performance: ~76.6% test accuracy on trained dataset",
    "Model artifact: Saved with joblib and loaded by Flask app at startup",
]
for item in ml_items:
    doc.add_paragraph(item, style="List Bullet")

# Authentication
doc.add_heading("Authentication", 2)
auth_items = [
    "Google OAuth: Verify tokens server-side (Google token verification) and issue JWTs",
    "Protected endpoints with @jwt_required for create/update/predict routes",
]
for item in auth_items:
    doc.add_paragraph(item, style="List Bullet")

# Chatbot & AI
doc.add_heading("Chatbot & AI", 2)
chatbot_items = [
    "Rule-based fallback for common commands (navigation, upload instructions, quick actions)",
    "Optional integration with Google Gemini (Generative Model) when VITE_GEMINI_API_KEY is set for richer conversational responses",
]
for item in chatbot_items:
    doc.add_paragraph(item, style="List Bullet")

# DevOps & Deployment
doc.add_heading("DevOps & Deployment", 2)
devops_items = [
    "Frontend: Netlify (build via npm run build, publish dist)",
    "Backend: Render (Gunicorn start command: gunicorn app:app)",
    "Environment variables for backend: DB URI, JWT secrets, Google client secrets",
    "CORS: Configure to allow both local dev (http://localhost:5173) and production Netlify domain",
]
for item in devops_items:
    doc.add_paragraph(item, style="List Bullet")

# ========== ML PIPELINE & DATA HANDLING ==========
doc.add_heading("ML Pipeline & Data Handling", 1)
ml_pipeline = [
    "Feature Engineering: Compute averageMarks from grade columns; compute attendance% from attendance/absent fields; derive feeStatus from tuition/debtor columns",
    "Naming Compatibility: Preprocessing includes fixes for naming mismatches (e.g., renames Nationality → Nacionality if model expects it)",
    "Serialization: Model probabilities (dict) are JSON-serialized before storing in SQLite to avoid binding errors",
    "Retraining: To update the model, re-run training with latest labeled data and save new joblib artifact",
]
for item in ml_pipeline:
    doc.add_paragraph(item, style="List Bullet")

# ========== PRODUCTION SETUP ==========
doc.add_heading("Production Setup Checklist", 1)
setup_items = [
    "Ensure backend has all required dependencies in requirements.txt (scikit-learn, pandas, flask, gunicorn, pymongo, etc.)",
    "For Render: use Gunicorn to run app:app",
    "Ensure DB migrations run to add new columns (class, department, feeStatus, prediction, confidence, probabilities)",
    "Set Netlify environment variables: VITE_API_URL, VITE_GOOGLE_CLIENT_ID, VITE_GEMINI_API_KEY",
    "Add Netlify URL to Google OAuth client authorized origins & redirect URIs",
    "Configure CORS in backend to allow Netlify domain",
]
for item in setup_items:
    doc.add_paragraph(item, style="List Bullet")

# ========== FUTURE ENHANCEMENTS ==========
doc.add_heading("Future Enhancements & Roadmap", 1)

doc.add_heading("Advanced AI & Analytics", 2)
ai_features = [
    "Emotion/Sentiment AI: Flag students with high stress or emotional distress using NLP for holistic well-being support",
    "Predictive Analytics 2.0: Deep Learning and Sequential Pattern Recognition for behavioral pattern mining and optimized dropout timing prediction",
    "Multimodal Data Integration: Analyze student engagement across digital learning platforms, library usage, and extracurricular participation",
]
for item in ai_features:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Intervention & Engagement Tools", 2)
intervention_features = [
    "Personalized Intervention Plans: Auto-generate tailored action plans with draft emails, parent communication templates, and remedial content",
    "Syllabus Simplifier: Convert complex topics into student-friendly lessons with AI-assisted content generation",
    "Financial Counseling Integration: Guide at-risk families to government schemes, scholarships, and fee waivers before crisis causes dropout",
    "Gamified Mentor & Student Engagement: Rewards, leaderboards, and recognition badges to encourage consistent mentorship and motivate students",
]
for item in intervention_features:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Accessibility & Scale", 2)
scale_features = [
    "Offline-First Design: PWA (Progressive Web App) allows dashboard to work offline and sync when connected — critical for low-infrastructure schools",
    "Multi-Language Support: Deliver support and communication in regional languages across literacy levels and rural regions",
    "Hybrid Parent Communication: Email, SMS, IVR support ensures accessibility regardless of digital literacy",
    "Macro Dashboard: National-level view for institutions to compare with regional and state-level performance, aligned with NEP 2020 and Digital India Vision 2030",
]
for item in scale_features:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Projected Impact", 2)
impact_points = [
    "30-40% reduction in secondary dropout within two academic years (especially for girls and SC/ST students in high-risk regions like Rajasthan)",
    "70% reduction in administrative time spent on dropout tracking",
    "Double mentor engagement and intervention rates",
    "Long-term socio-economic uplift: reduced early marriage, higher graduation rates, better family financial resilience",
]
for item in impact_points:
    doc.add_paragraph(item, style="List Bullet")

# ========== QUICK LOCAL RUN ==========
doc.add_heading("Quick Local Run (Developer)", 1)

doc.add_heading("Backend Setup", 2)
backend_code = """python -m venv .venv
.\.venv\Scripts\Activate.ps1
pip install -r backend/requirements.txt
cd backend
python app.py"""
p = doc.add_paragraph()
p.add_run(backend_code).font.name = "Consolas"
p.add_run(backend_code).font.size = Pt(10)

doc.add_heading("Frontend Setup", 2)
frontend_code = """npm install
npm run dev
# or build for production
npm run build"""
p = doc.add_paragraph()
p.add_run(frontend_code).font.name = "Consolas"
p.add_run(frontend_code).font.size = Pt(10)

doc.add_heading("Environment Variables", 2)
env_code = """VITE_API_URL=http://localhost:5000/api
VITE_GOOGLE_CLIENT_ID=<your-google-client-id>
VITE_GEMINI_API_KEY=<your-gemini-api-key>
JWT_SECRET_KEY=your_jwt_secret
MONGODB_URI=your_db_uri"""
p = doc.add_paragraph()
p.add_run(env_code).font.name = "Consolas"
p.add_run(env_code).font.size = Pt(10)

# ========== WHY MENTORAID MATTERS ==========
doc.add_heading("Why MentorAid Matters", 1)
closing = doc.add_paragraph(
    "MentorAid turns everyday student data — grades, attendance, fees, and demographics — into actionable "
    "early warnings. It helps teachers and counselors prioritize help, improves retention, and makes "
    "data-driven interventions repeatable and measurable. With lightweight deployment options "
    "(Netlify + Render) and straightforward CSV uploads, it fits into existing school workflows without "
    "heavy infrastructure."
)
closing.paragraph_format.space_after = Pt(12)

# ========== FOOTER ==========
doc.add_paragraph()
footer = doc.add_paragraph("—")
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact = doc.add_paragraph(
    "For more information, visit: https://mentoraid.netlify.app"
)
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.runs[0].font.size = Pt(10)
contact.runs[0].font.italic = True

# Save document
output_path = "d:/MentorAid/MentorAid-main/docs/MentorAid_Project_Overview.docx"
doc.save(output_path)
print(f"✅ Word document created successfully: {output_path}")
