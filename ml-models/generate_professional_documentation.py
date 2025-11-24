"""
Professional Documentation Generator for MentorAid Project
Generates comprehensive Word documentation with all model comparisons,
technology stack details, deep learning section, and professional structure.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()


def set_cell_background(cell, color):
    """Set background color for table cell"""
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def create_professional_documentation():
    """Create comprehensive professional documentation"""

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # ==================== COVER PAGE ====================
    # Title
    title = doc.add_heading("MentorAid", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(36)
    title_run.font.color.rgb = RGBColor(0, 102, 204)
    title_run.bold = True

    # Subtitle
    subtitle = doc.add_heading("AI-Powered Student Dropout Prediction System", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(20)
    subtitle_run.font.color.rgb = RGBColor(51, 51, 51)

    doc.add_paragraph("\n" * 3)

    # Project Details
    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details.add_run(
        "Comprehensive ML Model Analysis & Technical Documentation\n\n"
    ).bold = True
    details.add_run("Version: 1.0.0\n")
    details.add_run("Date: November 24, 2025\n")
    details.add_run("Status: Production Ready\n\n")
    details.add_run("Machine Learning • Deep Learning • Full Stack Development")

    add_page_break(doc)

    # ==================== TABLE OF CONTENTS ====================
    doc.add_heading("Table of Contents", level=1)

    toc_items = [
        ("1.", "Executive Summary", "4"),
        ("2.", "Technology Stack Overview", "5"),
        ("  2.1", "Frontend Technologies", "5"),
        ("  2.2", "Backend Technologies", "6"),
        ("  2.3", "Machine Learning Technologies", "7"),
        ("  2.4", "Development Tools", "8"),
        ("3.", "Dataset Overview & Analysis", "9"),
        ("4.", "Machine Learning Models", "10"),
        ("  4.1", "Traditional ML Models (Phase 1)", "10"),
        ("  4.2", "Enhanced ML Models (Phase 2)", "12"),
        ("  4.3", "Deep Learning Models", "14"),
        ("  4.4", "RF Optimization Models (Phase 3)", "16"),
        ("5.", "Comprehensive Model Comparisons", "18"),
        ("  5.1", "All Models Performance Matrix", "18"),
        ("  5.2", "Head-to-Head Comparisons", "20"),
        ("  5.3", "Category-wise Analysis", "25"),
        ("6.", "Best Model Deep Dive", "28"),
        ("7.", "Feature Engineering Analysis", "30"),
        ("8.", "System Architecture", "32"),
        ("  8.1", "Frontend Architecture (React + TypeScript)", "32"),
        ("  8.2", "Backend Architecture (Flask)", "34"),
        ("  8.3", "ML Pipeline Architecture", "36"),
        ("9.", "API Documentation", "38"),
        ("10.", "Deployment Guide", "40"),
        ("11.", "Model Strengths & Weaknesses", "42"),
        ("12.", "Recommendations & Future Work", "45"),
        ("13.", "Conclusion", "47"),
    ]

    toc_table = doc.add_table(rows=len(toc_items), cols=2)
    toc_table.style = "Light List"

    for idx, (num, title, page) in enumerate(toc_items):
        row = toc_table.rows[idx]
        row.cells[0].text = f"{num} {title}"
        row.cells[1].text = page
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    add_page_break(doc)

    # ==================== 1. EXECUTIVE SUMMARY ====================
    doc.add_heading("1. Executive Summary", level=1)

    exec_summary = doc.add_paragraph(
        "MentorAid is a comprehensive AI-powered student dropout prediction system designed to help "
        "educational institutions identify at-risk students early and implement timely interventions. "
        "This document provides a complete technical analysis of the machine learning pipeline, "
        "technology stack, and system architecture."
    )

    doc.add_heading("Key Achievements:", level=2)
    achievements = [
        "✅ Trained and evaluated 19 machine learning models across 3 development phases",
        "✅ Achieved 76.61% prediction accuracy with Random Forest + SMOTE + 20 engineered features",
        "✅ Implemented deep learning models using TensorFlow/Keras",
        "✅ Created production-ready REST API using Flask",
        "✅ Built modern responsive frontend using React, TypeScript, and Vite",
        "✅ Engineered 20 domain-specific features improving accuracy by 1.47%",
        "✅ Deployed complete full-stack application with automated startup scripts",
    ]

    for achievement in achievements:
        doc.add_paragraph(achievement, style="List Bullet")

    doc.add_heading("Project Overview:", level=2)

    # Project stats table
    stats_table = doc.add_table(rows=9, cols=2)
    stats_table.style = "Light Grid Accent 1"

    stats_data = [
        ("Total Models Trained", "19 models"),
        ("Best Model Accuracy", "76.61%"),
        ("Dataset Size", "4,424 students"),
        ("Features Used", "47 (28 original + 20 engineered - 7 dropped)"),
        ("Training Samples (after SMOTE)", "5,301"),
        ("Test Samples", "885"),
        ("Technology Stack", "React, TypeScript, Flask, scikit-learn, TensorFlow"),
        ("Development Phases", "3 (Initial, Enhanced, Optimized)"),
        ("Production Status", "Ready for Deployment ✅"),
    ]

    for idx, (key, value) in enumerate(stats_data):
        row = stats_table.rows[idx]
        row.cells[0].text = key
        row.cells[1].text = value
        set_cell_background(row.cells[0], "E7E6E6")
        row.cells[0].paragraphs[0].runs[0].bold = True

    add_page_break(doc)

    # ==================== 2. TECHNOLOGY STACK OVERVIEW ====================
    doc.add_heading("2. Technology Stack Overview", level=1)

    doc.add_paragraph(
        "MentorAid is built using modern, industry-standard technologies across the full stack. "
        "This section provides detailed explanations of each technology and its role in the system."
    )

    # 2.1 Frontend Technologies
    doc.add_heading("2.1 Frontend Technologies", level=2)

    doc.add_heading("React 18.x", level=3)
    doc.add_paragraph(
        "React is a powerful JavaScript library for building user interfaces, developed by Facebook. "
        "In MentorAid, React is used for:"
    )
    react_uses = [
        "Component-based architecture for reusable UI elements (Dashboard, StudentTable, FileUpload)",
        "Virtual DOM for efficient rendering and performance",
        "State management using React Hooks (useState, useEffect, useContext)",
        "Real-time UI updates when predictions are received from backend",
        "Conditional rendering for different user states (loading, error, success)",
    ]
    for use in react_uses:
        doc.add_paragraph(use, style="List Bullet 2")

    doc.add_heading("TypeScript 5.x", level=3)
    doc.add_paragraph(
        "TypeScript is a strongly-typed superset of JavaScript that adds static type checking. "
        "In MentorAid, TypeScript provides:"
    )
    ts_uses = [
        "Type safety for Student interface, preventing runtime errors",
        "IntelliSense and autocomplete in VS Code for faster development",
        "Early error detection during development (compile-time vs runtime)",
        "Better code documentation through type definitions",
        "Refactoring support with confidence",
    ]
    for use in ts_uses:
        doc.add_paragraph(use, style="List Bullet 2")

    doc.add_paragraph("\nExample TypeScript interface used in MentorAid:")
    code = doc.add_paragraph(
        "interface Student {\n"
        "  id: string;\n"
        "  name: string;\n"
        '  status: "Dropout" | "Enrolled" | "Graduate";\n'
        "  prediction?: string;\n"
        "  confidence?: number;\n"
        "  probabilities?: {\n"
        "    Dropout: number;\n"
        "    Enrolled: number;\n"
        "    Graduate: number;\n"
        "  };\n"
        "}"
    )
    code.style = "No Spacing"
    code_run = code.runs[0]
    code_run.font.name = "Consolas"
    code_run.font.size = Pt(9)

    doc.add_heading("Vite 5.x", level=3)
    doc.add_paragraph("Vite is a next-generation frontend build tool that provides:")
    vite_uses = [
        "Lightning-fast development server with Hot Module Replacement (HMR)",
        "Instant page updates when code changes (no full page reload)",
        "Optimized production builds using Rollup",
        "Native ESM support for faster cold starts",
        "Built-in TypeScript support without configuration",
    ]
    for use in vite_uses:
        doc.add_paragraph(use, style="List Bullet 2")

    doc.add_heading("TailwindCSS 3.x", level=3)
    doc.add_paragraph("TailwindCSS is a utility-first CSS framework used for:")
    tailwind_uses = [
        "Rapid UI development with pre-built utility classes",
        "Responsive design without writing custom media queries",
        "Dark mode support with class-based switching",
        "Consistent design system across all components",
        "Small production bundle size (only used classes included)",
        "Custom theme configuration for brand colors and spacing",
    ]
    for use in tailwind_uses:
        doc.add_paragraph(use, style="List Bullet 2")

    add_page_break(doc)

    # 2.2 Backend Technologies
    doc.add_heading("2.2 Backend Technologies", level=2)

    doc.add_heading("Flask 3.0.0", level=3)
    doc.add_paragraph(
        "Flask is a lightweight Python web framework used for building the REST API. "
        "In MentorAid, Flask handles:"
    )
    flask_uses = [
        "RESTful API endpoints for predictions (/api/predict, /api/predict/batch)",
        "CSV file uploads and processing",
        "Model loading and inference",
        "Feature engineering pipeline (20 custom features)",
        "Error handling and logging",
        "CORS (Cross-Origin Resource Sharing) for React frontend communication",
    ]
    for use in flask_uses:
        doc.add_paragraph(use, style="List Bullet 2")

    doc.add_heading("Flask-CORS 4.0.0", level=3)
    doc.add_paragraph(
        "Flask-CORS enables Cross-Origin Resource Sharing, allowing the React frontend "
        "(running on localhost:5173) to communicate with the Flask backend (localhost:5000). "
        "This is essential for development where frontend and backend run on different ports."
    )

    doc.add_heading("Python 3.10.18", level=3)
    doc.add_paragraph("Python serves as the primary language for:")
    python_uses = [
        "Machine learning model development and training",
        "Backend API implementation",
        "Data preprocessing and feature engineering",
        "Model serialization and deserialization (joblib)",
        "Scientific computing with NumPy and pandas",
    ]
    for use in python_uses:
        doc.add_paragraph(use, style="List Bullet 2")

    add_page_break(doc)

    # 2.3 Machine Learning Technologies
    doc.add_heading("2.3 Machine Learning Technologies", level=2)

    doc.add_heading("scikit-learn 1.4.0", level=3)
    doc.add_paragraph("scikit-learn is the primary ML library used for:")
    sklearn_uses = [
        "Traditional ML algorithms (Random Forest, Gradient Boosting, SVM, etc.)",
        "Data preprocessing (StandardScaler, LabelEncoder)",
        "Model evaluation metrics (accuracy, precision, recall, F1-score)",
        "Cross-validation for hyperparameter tuning",
        "Feature selection and importance analysis",
        "Train-test split functionality",
    ]
    for use in sklearn_uses:
        doc.add_paragraph(use, style="List Bullet 2")

    doc.add_heading("imbalanced-learn 0.12.0", level=3)
    doc.add_paragraph(
        "imbalanced-learn provides techniques for handling class imbalance:"
    )
    imblearn_uses = [
        "SMOTE (Synthetic Minority Over-sampling Technique) - increases minority class samples",
        "ADASYN (Adaptive Synthetic Sampling) - focuses on difficult-to-learn examples",
        "BorderlineSMOTE - generates samples near decision boundaries",
        "SMOTETomek - combines SMOTE with Tomek link removal",
        "Improved recall for minority classes (Dropout, Enrolled)",
    ]
    for use in imblearn_uses:
        doc.add_paragraph(use, style="List Bullet 2")

    doc.add_heading("TensorFlow 2.15.0 & Keras", level=3)
    doc.add_paragraph("TensorFlow/Keras is used for deep learning models:")
    tf_uses = [
        "Neural network architecture design (Sequential, Dense layers)",
        "Activation functions (ReLU, Softmax)",
        "Optimization algorithms (Adam, SGD)",
        "Dropout layers for regularization",
        "Batch normalization for stable training",
        "Early stopping to prevent overfitting",
        "GPU acceleration support (when available)",
    ]
    for use in tf_uses:
        doc.add_paragraph(use, style="List Bullet 2")

    doc.add_heading("XGBoost 2.0.0 & LightGBM 4.1.0", level=3)
    doc.add_paragraph("Gradient boosting libraries for high-performance models:")
    gb_uses = [
        "XGBoost - Extreme Gradient Boosting with regularization",
        "LightGBM - Light Gradient Boosting Machine (faster training)",
        "Handle missing values automatically",
        "Built-in feature importance calculation",
        "Support for custom evaluation metrics",
    ]
    for use in gb_uses:
        doc.add_paragraph(use, style="List Bullet 2")

    doc.add_heading("pandas 2.2.0 & NumPy 1.26.0", level=3)
    doc.add_paragraph("Data manipulation and numerical computing:")
    data_uses = [
        "pandas - CSV file loading, DataFrame operations, feature engineering",
        "NumPy - Array operations, mathematical computations, matrix operations",
        "Data cleaning (handling missing values, outliers)",
        "Statistical analysis and aggregations",
        "Data transformation and normalization",
    ]
    for use in data_uses:
        doc.add_paragraph(use, style="List Bullet 2")

    add_page_break(doc)

    # 2.4 Development Tools
    doc.add_heading("2.4 Development Tools & Infrastructure", level=2)

    doc.add_heading("Node.js & npm", level=3)
    doc.add_paragraph("Node.js runtime and npm package manager for:")
    node_uses = [
        "Frontend dependency management (package.json)",
        "Running Vite development server",
        "Build process for production deployment",
        "Script automation (npm run dev, npm run build)",
    ]
    for use in node_uses:
        doc.add_paragraph(use, style="List Bullet 2")

    doc.add_heading("Git & GitHub", level=3)
    doc.add_paragraph("Version control and collaboration:")
    git_uses = [
        "Source code versioning and history tracking",
        "Branch management (main, ml-integration)",
        "Collaboration with team members",
        "Code backup and disaster recovery",
        "Issue tracking and project management",
    ]
    for use in git_uses:
        doc.add_paragraph(use, style="List Bullet 2")

    doc.add_heading("PowerShell Scripts", level=3)
    doc.add_paragraph("Automated deployment scripts:")
    ps_uses = [
        "start-backend.ps1 - Automated backend setup and launch",
        "start-frontend.ps1 - Automated frontend setup and launch",
        "start-all.ps1 - One-click full stack deployment",
        "Virtual environment creation and activation",
        "Dependency installation automation",
    ]
    for use in ps_uses:
        doc.add_paragraph(use, style="List Bullet 2")

    doc.add_heading("VS Code", level=3)
    doc.add_paragraph("Primary development environment with:")
    vscode_uses = [
        "Integrated terminal for running scripts",
        "TypeScript IntelliSense and type checking",
        "Python extension for Jupyter notebooks",
        "Git integration for version control",
        "ESLint for code quality",
        "Prettier for code formatting",
    ]
    for use in vscode_uses:
        doc.add_paragraph(use, style="List Bullet 2")

    add_page_break(doc)

    # ==================== 3. DATASET OVERVIEW ====================
    doc.add_heading("3. Dataset Overview & Analysis", level=1)

    doc.add_paragraph(
        "The MentorAid system uses a comprehensive student dataset containing academic performance, "
        "demographic, and socioeconomic information."
    )

    doc.add_heading("Dataset Statistics:", level=2)

    dataset_table = doc.add_table(rows=11, cols=2)
    dataset_table.style = "Light Grid Accent 1"

    dataset_data = [
        ("Total Students", "4,424"),
        ("Original Features", "28"),
        ("Engineered Features", "20"),
        ("Total Features (Production)", "47"),
        ("Dropped Features", "7 (low correlation)"),
        ("Training Samples", "3,539 (80%)"),
        ("Test Samples", "885 (20%)"),
        ("Training After SMOTE", "5,301"),
        ("Target Classes", "3 (Dropout, Enrolled, Graduate)"),
        ("Class Distribution", "Dropout: 32%, Enrolled: 18%, Graduate: 50%"),
        ("Missing Values", "Handled during preprocessing"),
    ]

    for idx, (key, value) in enumerate(dataset_data):
        row = dataset_table.rows[idx]
        row.cells[0].text = key
        row.cells[1].text = value
        set_cell_background(row.cells[0], "E7E6E6")
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_heading("Original Features (28):", level=2)

    features_list = [
        "Marital Status - Student's marital status",
        "Application Mode - Method of application to university",
        "Application Order - Application preference order",
        "Course - Enrolled course",
        "Daytime/Evening Attendance - Class schedule type",
        "Previous Qualification - Educational background",
        "Nationality - Student nationality",
        "Mother's Qualification - Educational level of mother",
        "Father's Qualification - Educational level of father",
        "Mother's Occupation - Occupation type of mother",
        "Father's Occupation - Occupation type of father",
        "Displaced - Whether student is displaced",
        "Educational Special Needs - Special educational requirements",
        "Debtor - Outstanding payment status",
        "Tuition Fees Up to Date - Fee payment status",
        "Gender - Student gender",
        "Scholarship Holder - Scholarship status",
        "Age at Enrollment - Age when enrolled",
        "International - International student status",
        "Curricular Units 1st Sem (Credited) - Units credited in semester 1",
        "Curricular Units 1st Sem (Enrolled) - Units enrolled in semester 1",
        "Curricular Units 1st Sem (Evaluations) - Evaluations in semester 1",
        "Curricular Units 1st Sem (Approved) - Units approved in semester 1",
        "Curricular Units 1st Sem (Grade) - Average grade in semester 1",
        "Curricular Units 2nd Sem (Credited) - Units credited in semester 2",
        "Curricular Units 2nd Sem (Enrolled) - Units enrolled in semester 2",
        "Curricular Units 2nd Sem (Evaluations) - Evaluations in semester 2",
        "Curricular Units 2nd Sem (Approved) - Units approved in semester 2",
        "Curricular Units 2nd Sem (Grade) - Average grade in semester 2",
    ]

    for feature in features_list:
        doc.add_paragraph(feature, style="List Bullet")

    add_page_break(doc)

    # ==================== 4. MACHINE LEARNING MODELS ====================
    doc.add_heading("4. Machine Learning Models", level=1)

    doc.add_paragraph(
        "This section provides comprehensive analysis of all 19 models trained across 3 development phases, "
        "including traditional ML models, deep learning models, and optimized variants."
    )

    # 4.1 Traditional ML Models (Phase 1)
    doc.add_heading("4.1 Traditional ML Models (Phase 1)", level=2)

    doc.add_paragraph(
        "Phase 1 focused on establishing baseline performance using standard machine learning algorithms "
        "with minimal hyperparameter tuning. These models used the original 28 features without engineering."
    )

    # Phase 1 models table
    phase1_table = doc.add_table(rows=10, cols=4)
    phase1_table.style = "Light Grid Accent 1"

    # Header row
    header_cells = phase1_table.rows[0].cells
    header_cells[0].text = "Model"
    header_cells[1].text = "Accuracy"
    header_cells[2].text = "Training Time"
    header_cells[3].text = "Key Characteristics"
    for cell in header_cells:
        set_cell_background(cell, "4472C4")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].bold = True

    phase1_data = [
        ("Random Forest", "75.14%", "~2 sec", "Ensemble of 100 decision trees"),
        ("Gradient Boosting", "75.14%", "~5 sec", "Sequential boosting approach"),
        ("AdaBoost", "74.80%", "~3 sec", "Adaptive boosting"),
        ("Extra Trees", "74.12%", "~2 sec", "Randomized decision trees"),
        ("Logistic Regression", "73.90%", "~1 sec", "Linear classification"),
        ("Decision Tree", "71.64%", "~1 sec", "Single tree classifier"),
        ("K-Nearest Neighbors", "69.27%", "~1 sec", "Distance-based classification"),
        ("Support Vector Machine", "68.36%", "~10 sec", "Kernel-based separator"),
        ("Naive Bayes", "62.82%", "~1 sec", "Probabilistic classifier"),
    ]

    for idx, (model, acc, time, char) in enumerate(phase1_data, 1):
        row = phase1_table.rows[idx]
        row.cells[0].text = model
        row.cells[1].text = acc
        row.cells[2].text = time
        row.cells[3].text = char

    doc.add_heading("Phase 1 Key Findings:", level=3)
    phase1_findings = [
        "✓ Random Forest and Gradient Boosting tied for best performance (75.14%)",
        "✓ Ensemble methods outperformed single classifiers",
        "✓ Tree-based models showed strong performance (71-75%)",
        "✓ Naive Bayes struggled with feature dependencies (62.82%)",
        "✓ SVM was computationally expensive with limited benefit",
        "✓ Identified need for feature engineering to improve performance",
    ]
    for finding in phase1_findings:
        doc.add_paragraph(finding, style="List Bullet")

    add_page_break(doc)

    # 4.2 Enhanced ML Models (Phase 2)
    doc.add_heading("4.2 Enhanced ML Models (Phase 2)", level=2)

    doc.add_paragraph(
        "Phase 2 introduced feature engineering (20 new features), SMOTE class balancing, and hyperparameter "
        "tuning. This phase achieved the best overall performance."
    )

    # Phase 2 models table
    phase2_table = doc.add_table(rows=7, cols=5)
    phase2_table.style = "Light Grid Accent 1"

    # Header row
    header_cells = phase2_table.rows[0].cells
    header_cells[0].text = "Model"
    header_cells[1].text = "Accuracy"
    header_cells[2].text = "Precision"
    header_cells[3].text = "Recall"
    header_cells[4].text = "F1-Score"
    for cell in header_cells:
        set_cell_background(cell, "70AD47")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].bold = True

    phase2_data = [
        ("RF + SMOTE + 20 Features ⭐", "76.61%", "0.71", "0.71", "0.71"),
        ("Random Forest (Tuned)", "75.37%", "0.70", "0.70", "0.70"),
        ("XGBoost (Tuned)", "75.25%", "0.70", "0.70", "0.70"),
        ("Stacking Ensemble", "74.92%", "0.69", "0.69", "0.69"),
        ("LightGBM", "74.80%", "0.69", "0.69", "0.69"),
        ("XGBoost (Baseline)", "74.35%", "0.68", "0.68", "0.68"),
    ]

    for idx, (model, acc, prec, rec, f1) in enumerate(phase2_data, 1):
        row = phase2_table.rows[idx]
        row.cells[0].text = model
        row.cells[1].text = acc
        row.cells[2].text = prec
        row.cells[3].text = rec
        row.cells[4].text = f1
        if "⭐" in model:
            set_cell_background(row.cells[0], "FFF2CC")

    doc.add_heading("Phase 2 Breakthroughs:", level=3)
    phase2_findings = [
        "🏆 Best model achieved 76.61% accuracy (+1.47% from Phase 1)",
        "✓ Feature engineering created 20 domain-specific features",
        "✓ SMOTE balancing improved minority class recall",
        "✓ Hyperparameter tuning optimized n_estimators, max_depth, min_samples_split",
        "✓ Gradient boosting variants (XGBoost, LightGBM) showed competitive performance",
        "✓ Stacking ensemble combined multiple models but didn't exceed best individual model",
        "✓ Confirmed simpler models with good features > complex models with basic features",
    ]
    for finding in phase2_findings:
        doc.add_paragraph(finding, style="List Bullet")

    add_page_break(doc)

    # 4.3 Deep Learning Models
    doc.add_heading("4.3 Deep Learning Models", level=2)

    doc.add_paragraph(
        "Deep learning models were developed using TensorFlow/Keras to explore neural network architectures "
        "for this classification task. Multiple architectures were tested with varying complexity."
    )

    doc.add_heading("Neural Network Architectures Tested:", level=3)

    # Deep Learning table
    dl_table = doc.add_table(rows=5, cols=4)
    dl_table.style = "Light Grid Accent 1"

    # Header row
    header_cells = dl_table.rows[0].cells
    header_cells[0].text = "Architecture"
    header_cells[1].text = "Layers"
    header_cells[2].text = "Parameters"
    header_cells[3].text = "Performance"
    for cell in header_cells:
        set_cell_background(cell, "FFC000")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].bold = True

    dl_data = [
        ("Simple NN", "3 Dense (128, 64, 3)", "~12K", "72-73%"),
        ("Deep NN", "5 Dense (256, 128, 64, 32, 3)", "~45K", "73-74%"),
        ("NN + Dropout", "4 Dense + 2 Dropout", "~25K", "74-75%"),
        ("NN + Batch Norm", "4 Dense + 2 BatchNorm", "~28K", "74-75%"),
    ]

    for idx, (arch, layers, params, perf) in enumerate(dl_data, 1):
        row = dl_table.rows[idx]
        row.cells[0].text = arch
        row.cells[1].text = layers
        row.cells[2].text = params
        row.cells[3].text = perf

    doc.add_heading("Deep Learning Implementation Details:", level=3)

    dl_details = [
        "Activation Functions: ReLU (hidden layers), Softmax (output layer)",
        "Optimizer: Adam (learning_rate=0.001)",
        "Loss Function: Categorical Crossentropy",
        "Regularization: Dropout (0.3-0.5) and L2 regularization",
        "Batch Size: 32",
        "Epochs: 50-100 with Early Stopping (patience=10)",
        "Callbacks: ModelCheckpoint, EarlyStopping, ReduceLROnPlateau",
        "Validation Split: 20% of training data",
        "Data Normalization: StandardScaler applied before training",
    ]
    for detail in dl_details:
        doc.add_paragraph(detail, style="List Bullet")

    doc.add_heading("Deep Learning Results & Analysis:", level=3)

    doc.add_paragraph(
        "While deep learning models showed competitive performance (74-75%), they did not surpass "
        "the best traditional ML model (76.61%). Key observations:"
    )

    dl_observations = [
        "❌ Neural networks did not outperform Random Forest + SMOTE",
        "⚠ Limited training data (4,424 samples) insufficient for deep learning",
        "⚠ Overfitting occurred despite regularization (Dropout, L2)",
        "⚠ Longer training time (5-10 minutes vs <5 seconds for RF)",
        "⚠ Higher computational requirements (GPU beneficial but not necessary for RF)",
        "⚠ Less interpretable than tree-based models (black box)",
        "✓ Useful as baseline comparison to validate RF performance",
        "📊 Deep learning typically requires 20,000+ samples to excel",
    ]
    for obs in dl_observations:
        doc.add_paragraph(obs, style="List Bullet")

    doc.add_heading("Why Traditional ML Won:", level=3)

    doc.add_paragraph("For this specific dataset and problem:")

    ml_advantages = [
        "✅ Tabular Data: Traditional ML (RF, GBM) excel with structured tabular data",
        "✅ Small Dataset: 4,424 samples insufficient for deep learning to learn complex patterns",
        "✅ Feature Engineering: Carefully crafted features captured domain knowledge",
        "✅ Interpretability: Feature importance from RF helps understand predictions",
        "✅ Speed: RF trains in seconds vs minutes for neural networks",
        "✅ Simplicity: Fewer hyperparameters to tune",
        "✅ No GPU Required: Can run on any machine",
        "✅ Production Ready: Easier to deploy and maintain",
    ]
    for adv in ml_advantages:
        doc.add_paragraph(adv, style="List Bullet")

    add_page_break(doc)

    # 4.4 RF Optimization (Phase 3)
    doc.add_heading("4.4 RF Optimization Models (Phase 3)", level=2)

    doc.add_paragraph(
        "Phase 3 attempted aggressive optimization of Random Forest through extensive feature engineering "
        "(45+ features), multiple balancing techniques, and intensive hyperparameter tuning."
    )

    # Phase 3 models table
    phase3_table = doc.add_table(rows=5, cols=5)
    phase3_table.style = "Light Grid Accent 1"

    # Header row
    header_cells = phase3_table.rows[0].cells
    header_cells[0].text = "Model Variant"
    header_cells[1].text = "Accuracy"
    header_cells[2].text = "CV Score"
    header_cells[3].text = "Training Time"
    header_cells[4].text = "Result"
    for cell in header_cells:
        set_cell_background(cell, "C55A11")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].bold = True

    phase3_data = [
        ("RF Top 40 Features", "75.48%", "N/A", "~3 sec", "Below baseline"),
        ("RF Tuned (200 iter)", "74.69%", "85.18%", "5 hours", "Overfitted"),
        ("ExtraTrees", "74.58%", "N/A", "~2 sec", "Below baseline"),
        ("RF GridSearch", "74.01%", "85.53%", "2 hours", "Severe overfit"),
    ]

    for idx, (model, acc, cv, time, result) in enumerate(phase3_data, 1):
        row = phase3_table.rows[idx]
        row.cells[0].text = model
        row.cells[1].text = acc
        row.cells[2].text = cv
        row.cells[3].text = time
        row.cells[4].text = result
        set_cell_background(row.cells[4], "FFE699")

    doc.add_heading("Phase 3 Critical Learnings:", level=3)

    phase3_learnings = [
        "❌ ALL optimization attempts performed WORSE than Phase 2 baseline (76.61%)",
        "⚠ 45+ engineered features added noise instead of signal",
        "⚠ 200-iteration RandomizedSearchCV: 85.18% CV but only 74.69% test (10% overfit gap)",
        "⚠ GridSearch fine-tuning: 85.53% CV but only 74.01% test (11% overfit gap)",
        "⚠ 7+ hours of training time yielded worse results",
        "⚠ Feature selection (top 40) still below baseline (75.48% vs 76.61%)",
        "✓ Proved simpler is better for this dataset size",
        "✓ Validated Phase 2 model as optimal configuration",
        "📊 Key Insight: Dataset too small (4,424) for aggressive optimization",
    ]
    for learning in phase3_learnings:
        doc.add_paragraph(learning, style="List Bullet")

    add_page_break(doc)

    # ==================== 5. COMPREHENSIVE MODEL COMPARISONS ====================
    doc.add_heading("5. Comprehensive Model Comparisons", level=1)

    doc.add_paragraph(
        "This section provides exhaustive comparisons between all 19 models trained, analyzing "
        "performance, strengths, weaknesses, and use case recommendations."
    )

    # 5.1 All Models Performance Matrix
    doc.add_heading("5.1 All Models Performance Matrix", level=2)

    doc.add_paragraph("Complete ranking of all 19 models:")

    # Complete ranking table
    ranking_table = doc.add_table(rows=20, cols=6)
    ranking_table.style = "Medium Grid 1 Accent 1"

    # Header row
    header_cells = ranking_table.rows[0].cells
    headers = ["Rank", "Model", "Accuracy", "Phase", "Training Time", "Status"]
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        set_cell_background(header_cells[idx], "2E75B6")
        header_cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        header_cells[idx].paragraphs[0].runs[0].bold = True

    ranking_data = [
        (
            "🥇 1",
            "RF + SMOTE + 20 Features",
            "76.61%",
            "Phase 2",
            "~3 sec",
            "✅ PRODUCTION",
        ),
        ("2", "Random Forest", "75.14%", "Phase 1", "~2 sec", "Baseline"),
        ("2", "Gradient Boosting", "75.14%", "Phase 1", "~5 sec", "Baseline"),
        ("4", "RF Top 40 Features", "75.48%", "Phase 3", "~3 sec", "Optimization"),
        ("5", "RF Tuned", "75.37%", "Phase 2", "~3 sec", "Enhanced"),
        ("6", "XGBoost Tuned", "75.25%", "Phase 2", "~4 sec", "Enhanced"),
        ("7", "Stacking Ensemble", "74.92%", "Phase 2", "~10 sec", "Enhanced"),
        ("8", "AdaBoost", "74.80%", "Phase 1", "~3 sec", "Baseline"),
        ("9", "LightGBM", "74.80%", "Phase 2", "~2 sec", "Enhanced"),
        ("10", "RF Tuned (200 iter)", "74.69%", "Phase 3", "5 hours", "Overfitted"),
        ("11", "ExtraTrees", "74.58%", "Phase 3", "~2 sec", "Optimization"),
        ("12", "XGBoost Baseline", "74.35%", "Phase 2", "~3 sec", "Enhanced"),
        ("13", "Extra Trees", "74.12%", "Phase 1", "~2 sec", "Baseline"),
        ("14", "RF GridSearch", "74.01%", "Phase 3", "2 hours", "Overfitted"),
        ("15", "Logistic Regression", "73.90%", "Phase 1", "~1 sec", "Baseline"),
        ("16", "Decision Tree", "71.64%", "Phase 1", "~1 sec", "Baseline"),
        ("17", "K-Nearest Neighbors", "69.27%", "Phase 1", "~1 sec", "Baseline"),
        ("18", "Support Vector Machine", "68.36%", "Phase 1", "~10 sec", "Baseline"),
        ("19", "Naive Bayes", "62.82%", "Phase 1", "~1 sec", "Baseline"),
    ]

    for idx, data in enumerate(ranking_data, 1):
        row = ranking_table.rows[idx]
        for col_idx, value in enumerate(data):
            row.cells[col_idx].text = value

        if idx == 1:  # Highlight best model
            for cell in row.cells:
                set_cell_background(cell, "C6EFCE")

    add_page_break(doc)

    # 5.2 Head-to-Head Comparisons
    doc.add_heading("5.2 Head-to-Head Model Comparisons", level=2)

    doc.add_paragraph(
        "Detailed comparisons between models to understand relative strengths and weaknesses:"
    )

    # Comparison 1: RF+SMOTE vs Random Forest
    doc.add_heading("RF + SMOTE + 20 Features vs. Random Forest (Baseline)", level=3)

    comparison1 = [
        ("Metric", "RF + SMOTE + Features", "Random Forest", "Difference"),
        ("Accuracy", "76.61%", "75.14%", "+1.47%"),
        ("Dropout Precision", "0.82", "0.78", "+0.04"),
        ("Dropout Recall", "0.71", "0.65", "+0.06"),
        ("Enrolled Precision", "0.47", "0.42", "+0.05"),
        ("Enrolled Recall", "0.53", "0.48", "+0.05"),
        ("Graduate Precision", "0.83", "0.82", "+0.01"),
        ("Graduate Recall", "0.86", "0.85", "+0.01"),
        ("Training Time", "~3 sec", "~2 sec", "+1 sec"),
        ("Features", "47", "28", "+19"),
    ]

    comp1_table = doc.add_table(rows=len(comparison1), cols=4)
    comp1_table.style = "Light Grid Accent 1"

    for idx, row_data in enumerate(comparison1):
        row = comp1_table.rows[idx]
        for col_idx, value in enumerate(row_data):
            row.cells[col_idx].text = value
        if idx == 0:
            for cell in row.cells:
                set_cell_background(cell, "5B9BD5")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph("\nKey Takeaways:")
    comp1_takeaways = [
        "✓ Feature engineering provided 1.47% accuracy boost",
        "✓ Significant improvement in minority class recall (Dropout +6%, Enrolled +5%)",
        "✓ SMOTE balancing addressed class imbalance effectively",
        "✓ Minimal increase in training time (+1 second)",
        "✓ Graduate class already well-predicted, minimal improvement needed",
    ]
    for takeaway in comp1_takeaways:
        doc.add_paragraph(takeaway, style="List Bullet")

    # Comparison 2: RF vs Gradient Boosting
    doc.add_heading("Random Forest vs. Gradient Boosting", level=3)

    comparison2 = [
        ("Metric", "Random Forest", "Gradient Boosting", "Winner"),
        ("Accuracy", "75.14%", "75.14%", "Tie"),
        ("Training Time", "~2 sec", "~5 sec", "RF (faster)"),
        ("Interpretability", "High", "Medium", "RF"),
        ("Overfitting Risk", "Low", "Medium", "RF"),
        ("Parallelization", "Excellent", "Sequential", "RF"),
        ("Hyperparameter Tuning", "Easier", "More complex", "RF"),
    ]

    comp2_table = doc.add_table(rows=len(comparison2), cols=4)
    comp2_table.style = "Light Grid Accent 1"

    for idx, row_data in enumerate(comparison2):
        row = comp2_table.rows[idx]
        for col_idx, value in enumerate(row_data):
            row.cells[col_idx].text = value
        if idx == 0:
            for cell in row.cells:
                set_cell_background(cell, "70AD47")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph("\nVerdict: Random Forest preferred due to speed and simplicity")

    # Comparison 3: XGBoost vs LightGBM vs RF
    doc.add_heading("Gradient Boosting Variants: XGBoost vs LightGBM vs RF", level=3)

    comparison3 = [
        ("Metric", "RF + SMOTE", "XGBoost", "LightGBM"),
        ("Accuracy", "76.61%", "75.25%", "74.80%"),
        ("Training Time", "~3 sec", "~4 sec", "~2 sec"),
        ("Memory Usage", "Low", "Medium", "Low"),
        ("GPU Support", "No", "Yes", "Yes"),
        ("Overfitting Control", "Excellent", "Good", "Good"),
        ("Production Readiness", "Excellent", "Good", "Good"),
    ]

    comp3_table = doc.add_table(rows=len(comparison3), cols=4)
    comp3_table.style = "Light Grid Accent 1"

    for idx, row_data in enumerate(comparison3):
        row = comp3_table.rows[idx]
        for col_idx, value in enumerate(row_data):
            row.cells[col_idx].text = value
        if idx == 0:
            for cell in row.cells:
                set_cell_background(cell, "FFC000")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph(
        "\nVerdict: Random Forest wins on accuracy, simplicity, and production readiness"
    )

    # Comparison 4: Traditional ML vs Deep Learning
    doc.add_heading("Traditional ML vs Deep Learning", level=3)

    comparison4 = [
        ("Aspect", "RF + SMOTE", "Neural Networks", "Winner"),
        ("Best Accuracy", "76.61%", "74-75%", "RF (+1.6%)"),
        ("Training Time", "~3 sec", "5-10 min", "RF"),
        ("Data Requirements", "~5K samples", "20K+ samples", "RF"),
        ("Interpretability", "High (feature importance)", "Low (black box)", "RF"),
        ("Hardware Requirements", "Any CPU", "GPU beneficial", "RF"),
        ("Deployment Complexity", "Simple (pickle)", "Complex (TF serving)", "RF"),
        ("Maintenance", "Easy", "Moderate", "RF"),
        ("Overfitting Risk", "Low", "High (small data)", "RF"),
    ]

    comp4_table = doc.add_table(rows=len(comparison4), cols=4)
    comp4_table.style = "Light Grid Accent 1"

    for idx, row_data in enumerate(comparison4):
        row = comp4_table.rows[idx]
        for col_idx, value in enumerate(row_data):
            row.cells[col_idx].text = value
        if idx == 0:
            for cell in row.cells:
                set_cell_background(cell, "ED7D31")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph(
        "\nConclusion: Random Forest is the clear winner for this dataset and use case"
    )

    add_page_break(doc)

    # 5.3 Category-wise Analysis
    doc.add_heading("5.3 Category-wise Performance Analysis", level=2)

    doc.add_heading("Ensemble Methods Performance", level=3)

    ensemble_data = [
        ("Model", "Accuracy", "Ensemble Type", "Complexity"),
        ("RF + SMOTE", "76.61%", "Bagging", "Medium"),
        ("Random Forest", "75.14%", "Bagging", "Low"),
        ("Gradient Boosting", "75.14%", "Boosting", "Medium"),
        ("Stacking Ensemble", "74.92%", "Stacking", "High"),
        ("AdaBoost", "74.80%", "Boosting", "Medium"),
        ("Extra Trees", "74.12%", "Bagging", "Low"),
    ]

    ensemble_table = doc.add_table(rows=len(ensemble_data), cols=4)
    ensemble_table.style = "Light Grid Accent 1"

    for idx, row_data in enumerate(ensemble_data):
        row = ensemble_table.rows[idx]
        for col_idx, value in enumerate(row_data):
            row.cells[col_idx].text = value
        if idx == 0:
            for cell in row.cells:
                set_cell_background(cell, "4472C4")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph("\nInsights:")
    ensemble_insights = [
        "✓ Bagging (RF) outperformed boosting (GB, AdaBoost)",
        "✓ Stacking added complexity without accuracy gain",
        "✓ Simple Random Forest competitive with complex ensembles",
    ]
    for insight in ensemble_insights:
        doc.add_paragraph(insight, style="List Bullet")

    doc.add_heading("Linear vs Non-Linear Models", level=3)

    linear_data = [
        ("Model Type", "Model", "Accuracy", "Decision Boundary"),
        ("Non-Linear", "RF + SMOTE", "76.61%", "Complex"),
        ("Non-Linear", "Random Forest", "75.14%", "Complex"),
        ("Non-Linear", "Gradient Boosting", "75.14%", "Complex"),
        ("Linear", "Logistic Regression", "73.90%", "Linear"),
        ("Non-Linear", "SVM (RBF)", "68.36%", "Non-linear"),
        ("Linear", "Naive Bayes", "62.82%", "Linear"),
    ]

    linear_table = doc.add_table(rows=len(linear_data), cols=4)
    linear_table.style = "Light Grid Accent 1"

    for idx, row_data in enumerate(linear_data):
        row = linear_table.rows[idx]
        for col_idx, value in enumerate(row_data):
            row.cells[col_idx].text = value
        if idx == 0:
            for cell in row.cells:
                set_cell_background(cell, "70AD47")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph(
        "\nConclusion: Non-linear models essential for this problem (complex feature interactions)"
    )

    add_page_break(doc)

    # ==================== REST OF DOCUMENT CONTINUES ====================
    # (Truncated for length - continuing with remaining sections)

    doc.add_heading("6. Best Model Deep Dive", level=1)
    doc.add_paragraph(
        "Comprehensive analysis of the production model: Random Forest + SMOTE + 20 Engineered Features"
    )

    doc.add_heading("Model Configuration:", level=2)
    config_data = [
        ("Parameter", "Value", "Explanation"),
        ("n_estimators", "300", "Number of decision trees in forest"),
        ("max_depth", "20", "Maximum depth of each tree"),
        ("min_samples_split", "10", "Minimum samples required to split node"),
        ("max_features", "sqrt", "Features considered at each split"),
        ("random_state", "42", "Reproducibility seed"),
        ("class_weight", "balanced", "Handle class imbalance"),
        ("Balancing Method", "SMOTE", "Synthetic minority oversampling"),
        ("Features", "47", "28 original + 20 engineered - 7 dropped"),
    ]

    config_table = doc.add_table(rows=len(config_data), cols=3)
    config_table.style = "Light Grid Accent 1"

    for idx, row_data in enumerate(config_data):
        row = config_table.rows[idx]
        for col_idx, value in enumerate(row_data):
            row.cells[col_idx].text = value
        if idx == 0:
            for cell in row.cells:
                set_cell_background(cell, "5B9BD5")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].bold = True

    doc.add_heading("Performance Metrics:", level=2)

    perf_data = [
        ("Class", "Precision", "Recall", "F1-Score", "Support"),
        ("Dropout", "0.82", "0.71", "0.76", "283"),
        ("Enrolled", "0.47", "0.53", "0.50", "160"),
        ("Graduate", "0.83", "0.86", "0.84", "442"),
        ("Weighted Avg", "0.77", "0.77", "0.77", "885"),
    ]

    perf_table = doc.add_table(rows=len(perf_data), cols=5)
    perf_table.style = "Light Grid Accent 1"

    for idx, row_data in enumerate(perf_data):
        row = perf_table.rows[idx]
        for col_idx, value in enumerate(row_data):
            row.cells[col_idx].text = value
        if idx == 0:
            for cell in row.cells:
                set_cell_background(cell, "70AD47")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].bold = True

    doc.add_heading("Top 10 Most Important Features:", level=2)

    features_data = [
        ("Rank", "Feature", "Importance", "Category"),
        ("1", "Curricular units 2nd sem (approved)", "6.86%", "Academic"),
        ("2", "completion_rate", "6.58%", "Engineered"),
        ("3", "total_approved", "6.21%", "Engineered"),
        ("4", "avg_approved", "5.82%", "Engineered"),
        ("5", "avg_grade", "4.70%", "Engineered"),
        ("6", "Curricular units 1st sem (approved)", "4.65%", "Academic"),
        ("7", "Curricular units 2nd sem (grade)", "4.32%", "Academic"),
        ("8", "total_evaluations", "3.98%", "Engineered"),
        ("9", "Curricular units 1st sem (grade)", "3.87%", "Academic"),
        ("10", "total_failure_rate", "3.65%", "Engineered"),
    ]

    features_table = doc.add_table(rows=len(features_data), cols=4)
    features_table.style = "Light Grid Accent 1"

    for idx, row_data in enumerate(features_data):
        row = features_table.rows[idx]
        for col_idx, value in enumerate(row_data):
            row.cells[col_idx].text = value
        if idx == 0:
            for cell in row.cells:
                set_cell_background(cell, "FFC000")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph("\nKey Insight: 6 of top 10 features are engineered (60%)")

    add_page_break(doc)

    # ==================== 7. FEATURE ENGINEERING ANALYSIS ====================
    doc.add_heading("7. Feature Engineering Analysis", level=1)

    doc.add_paragraph(
        "Feature engineering was crucial to achieving the best model performance (76.61%). "
        "This section details all 20 engineered features and their impact on prediction accuracy."
    )

    doc.add_heading("7.1 All 20 Engineered Features", level=2)

    # Feature categories
    doc.add_heading("Academic Performance Features (8 features):", level=3)

    academic_features = [
        (
            "avg_approved",
            "Average approved units across both semesters",
            "(sem1_approved + sem2_approved) / 2",
        ),
        (
            "total_approved",
            "Total units approved across both semesters",
            "sem1_approved + sem2_approved",
        ),
        (
            "total_evaluations",
            "Total evaluations across both semesters",
            "sem1_evaluations + sem2_evaluations",
        ),
        (
            "completion_rate",
            "Ratio of approved units to total evaluations",
            "total_approved / total_evaluations",
        ),
        (
            "avg_grade",
            "Average grade across both semesters",
            "(sem1_grade + sem2_grade) / 2",
        ),
        (
            "grade_improvement",
            "Grade improvement from semester 1 to 2",
            "sem2_grade - sem1_grade",
        ),
        (
            "approval_improvement",
            "Improvement in approved units",
            "sem2_approved - sem1_approved",
        ),
        ("evaluation_load", "Total academic workload", "sem1_enrolled + sem2_enrolled"),
    ]

    for idx, (name, desc, formula) in enumerate(academic_features, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"{name}").bold = True
        p.add_run(f" - {desc}\n   Formula: {formula}")

    doc.add_heading("Semester-wise Performance Features (6 features):", level=3)

    semester_features = [
        (
            "failure_rate_sem1",
            "Failure rate in semester 1",
            "(enrolled - approved) / enrolled for sem1",
        ),
        (
            "failure_rate_sem2",
            "Failure rate in semester 2",
            "(enrolled - approved) / enrolled for sem2",
        ),
        (
            "total_failure_rate",
            "Average failure rate across semesters",
            "(failure_rate_sem1 + failure_rate_sem2) / 2",
        ),
        (
            "sem1_performance",
            "Overall semester 1 performance score",
            "approved * grade for sem1",
        ),
        (
            "sem2_performance",
            "Overall semester 2 performance score",
            "approved * grade for sem2",
        ),
        (
            "performance_trend",
            "Performance trajectory",
            "sem2_performance - sem1_performance",
        ),
    ]

    for idx, (name, desc, formula) in enumerate(semester_features, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"{name}").bold = True
        p.add_run(f" - {desc}\n   Formula: {formula}")

    doc.add_heading("Credit Management Features (3 features):", level=3)

    credit_features = [
        ("total_credited", "Total credited units", "sem1_credited + sem2_credited"),
        (
            "credit_ratio",
            "Ratio of credited to enrolled units",
            "total_credited / evaluation_load",
        ),
        (
            "units_without_eval",
            "Units enrolled but not evaluated",
            "enrolled - evaluations",
        ),
    ]

    for idx, (name, desc, formula) in enumerate(credit_features, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"{name}").bold = True
        p.add_run(f" - {desc}\n   Formula: {formula}")

    doc.add_heading("Advanced Performance Metrics (3 features):", level=3)

    advanced_features = [
        (
            "weighted_grade",
            "Grade weighted by approved units",
            "avg_grade * total_approved",
        ),
        (
            "consistency_score",
            "Consistency between semesters",
            "1 - abs(sem1_grade - sem2_grade) / max(sem1_grade, sem2_grade)",
        ),
        (
            "academic_momentum",
            "Academic progress momentum",
            "(sem2_approved - sem1_approved) + (sem2_grade - sem1_grade)",
        ),
    ]

    for idx, (name, desc, formula) in enumerate(advanced_features, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"{name}").bold = True
        p.add_run(f" - {desc}\n   Formula: {formula}")

    doc.add_heading("7.2 Feature Engineering Impact", level=2)

    impact_data = [
        ("Metric", "Before Engineering", "After Engineering", "Improvement"),
        ("Model Accuracy", "75.14%", "76.61%", "+1.47%"),
        ("Dropout Recall", "0.65", "0.71", "+0.06 (+9.2%)"),
        ("Enrolled Recall", "0.48", "0.53", "+0.05 (+10.4%)"),
        ("Total Features", "28", "47", "+19 (+67.9%)"),
        ("Feature Importance (Top 10)", "4 original", "6 engineered", "60% engineered"),
    ]

    impact_table = doc.add_table(rows=len(impact_data), cols=4)
    impact_table.style = "Medium Grid 1 Accent 1"

    for idx, row_data in enumerate(impact_data):
        row = impact_table.rows[idx]
        for col_idx, value in enumerate(row_data):
            row.cells[col_idx].text = value
        if idx == 0:
            for cell in row.cells:
                set_cell_background(cell, "4472C4")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].bold = True

    doc.add_heading("7.3 Key Insights from Feature Engineering", level=2)

    fe_insights = [
        "✅ 1.47% accuracy improvement may seem small, but represents ~13 more students correctly classified out of 885",
        "✅ Engineered features dominate top importance (6 out of top 10 features)",
        "✅ completion_rate (#2 importance) captures critical academic success metric",
        "✅ Performance trend features help identify students losing momentum",
        "✅ Grade improvement features reveal recovery or decline patterns",
        "✅ Total failure rate aggregates risk across both semesters",
        "✅ Minimal computational overhead (features computed in milliseconds)",
        "✅ Domain knowledge embedded: academic success = approval rate + grades",
    ]

    for insight in fe_insights:
        doc.add_paragraph(insight, style="List Bullet")

    add_page_break(doc)

    # ==================== 8. SYSTEM ARCHITECTURE ====================
    doc.add_heading("8. System Architecture", level=1)

    doc.add_paragraph(
        "MentorAid is built as a modern three-tier architecture: React frontend, Flask REST API backend, "
        "and scikit-learn ML pipeline. This section details each component and their interactions."
    )

    doc.add_heading("8.1 Frontend Architecture (React + TypeScript)", level=2)

    doc.add_paragraph("Technology Stack:")
    frontend_stack = [
        "React 18.x - UI component library with Virtual DOM",
        "TypeScript 5.x - Static typing and type safety",
        "Vite 5.x - Build tool and development server",
        "TailwindCSS 3.x - Utility-first CSS framework",
        "React Router - Client-side routing",
        "React Context API - State management",
    ]
    for tech in frontend_stack:
        doc.add_paragraph(tech, style="List Bullet 2")

    doc.add_heading("Component Hierarchy:", level=3)

    hierarchy = [
        "App.tsx (Root)",
        "  ├── ThemeContext (Dark/Light mode)",
        "  ├── AuthContext (Authentication state)",
        "  └── Router",
        "      ├── Login.tsx",
        "      ├── Introduction.tsx",
        "      ├── Dashboard.tsx",
        "      │   ├── DashboardStats.tsx (Summary cards)",
        "      │   ├── FileUpload.tsx (CSV upload + API call)",
        "      │   ├── StudentTable.tsx (Results display)",
        "      │   └── ChartsSection.tsx (Visualizations)",
        "      └── StudentDetails.tsx",
        "          ├── StudentMetrics.tsx",
        "          ├── ProgressChart.tsx",
        "          ├── AIInsights.tsx",
        "          └── InterventionHistory.tsx",
    ]

    code_para = doc.add_paragraph()
    code_run = code_para.add_run("\n".join(hierarchy))
    code_run.font.name = "Consolas"
    code_run.font.size = Pt(9)

    doc.add_heading("Key Frontend Features:", level=3)

    frontend_features = [
        "CSV File Upload: Drag-and-drop interface with validation",
        "Real-time Predictions: Async API calls with loading states",
        "Type Safety: TypeScript interfaces for Student, Prediction, API responses",
        "State Management: React Context for global state (auth, theme)",
        "Error Handling: Try-catch blocks with user-friendly error messages",
        "Responsive Design: Mobile-first approach using TailwindCSS",
        "Dark Mode: Theme switching with persistent localStorage",
        "Optimized Rendering: React.memo for expensive components",
    ]

    for feature in frontend_features:
        doc.add_paragraph(feature, style="List Bullet")

    doc.add_heading("Data Flow (Frontend):", level=3)

    flow_steps = [
        "1. User uploads CSV file via FileUpload.tsx",
        "2. File validated (size, format, required columns)",
        "3. FormData created and sent to Flask backend via fetch()",
        "4. Loading spinner displayed during API call",
        "5. Response received with predictions array",
        "6. Predictions converted to Student[] objects",
        "7. Dashboard.tsx state updated",
        "8. StudentTable.tsx re-renders with new data",
        "9. DashboardStats.tsx shows summary (dropout %, graduate %, etc.)",
        "10. Toast notification confirms success",
    ]

    for step in flow_steps:
        doc.add_paragraph(step, style="List Number")

    add_page_break(doc)

    doc.add_heading("8.2 Backend Architecture (Flask)", level=2)

    doc.add_paragraph("Technology Stack:")
    backend_stack = [
        "Flask 3.0.0 - Micro web framework",
        "Flask-CORS 4.0.0 - Cross-origin resource sharing",
        "Python 3.10.18 - Core language",
        "pandas 2.2.0 - Data manipulation",
        "NumPy 1.26.0 - Numerical computing",
        "scikit-learn 1.4.0 - ML model loading",
        "joblib 1.3.2 - Model serialization",
    ]
    for tech in backend_stack:
        doc.add_paragraph(tech, style="List Bullet 2")

    doc.add_heading("Backend Components:", level=3)

    backend_components = [
        "app.py - Main Flask application",
        "  ├── Model Loading - Loads pkl files on startup",
        "  ├── Feature Engineering - engineer_features() function",
        "  ├── Preprocessing - preprocess_data() function",
        "  ├── Prediction Pipeline - Full ML inference",
        "  └── API Routes - 5 RESTful endpoints",
    ]

    code_para2 = doc.add_paragraph()
    code_run2 = code_para2.add_run("\n".join(backend_components))
    code_run2.font.name = "Consolas"
    code_run2.font.size = Pt(9)

    doc.add_heading("Prediction Pipeline:", level=3)

    pipeline_steps = [
        "1. Receive CSV file from frontend",
        "2. Load CSV into pandas DataFrame",
        "3. Validate required columns (28 original features)",
        "4. Engineer 20 additional features",
        "5. Drop 7 low-correlation features",
        "6. Scale features using StandardScaler",
        "7. Predict using Random Forest model",
        "8. Get prediction probabilities",
        "9. Calculate confidence scores",
        "10. Assign risk levels (High/Medium/Low)",
        "11. Format response as JSON",
        "12. Return predictions + summary statistics",
    ]

    for step in pipeline_steps:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("8.3 ML Pipeline Architecture", level=2)

    doc.add_paragraph("Model Artifacts:")

    artifacts = [
        "random_forest_model.pkl - Trained Random Forest classifier (300 trees)",
        "scaler.pkl - StandardScaler fitted on training data",
        "label_encoder.pkl - LabelEncoder for class names",
        "feature_names.pkl - List of 47 feature names in correct order",
        "model_metadata.json - Model configuration and performance metrics",
        "feature_importance.csv - Ranked feature importance scores",
    ]

    for artifact in artifacts:
        doc.add_paragraph(artifact, style="List Bullet")

    doc.add_heading("Complete System Data Flow:", level=3)

    system_flow = [
        "User (Browser) → React Frontend (localhost:5173)",
        "  ↓ [CSV Upload]",
        "  POST /api/predict/batch",
        "  ↓",
        "Flask Backend (localhost:5000)",
        "  ↓ [Load CSV]",
        "  pandas.read_csv()",
        "  ↓ [Engineer Features]",
        "  20 new features created",
        "  ↓ [Scale Features]",
        "  StandardScaler.transform()",
        "  ↓ [Predict]",
        "  RandomForest.predict() + predict_proba()",
        "  ↓ [Format Response]",
        "  JSON: {predictions[], summary{}}",
        "  ↓",
        "React Frontend",
        "  ↓ [Update State]",
        "  Dashboard displays results",
        "  ↓",
        "User sees predictions in table",
    ]

    flow_para = doc.add_paragraph()
    flow_run = flow_para.add_run("\n".join(system_flow))
    flow_run.font.name = "Consolas"
    flow_run.font.size = Pt(9)

    add_page_break(doc)

    # ==================== 9. API DOCUMENTATION ====================
    doc.add_heading("9. API Documentation", level=1)

    doc.add_paragraph(
        "The Flask backend exposes 5 RESTful API endpoints for health checks, model info, "
        "and predictions. All endpoints return JSON responses."
    )

    # Endpoint 1: Health Check
    doc.add_heading("9.1 GET /api/health", level=2)
    doc.add_paragraph("Health check endpoint to verify backend is running.")

    doc.add_heading("Request:", level=3)
    doc.add_paragraph(
        "Method: GET\nURL: http://localhost:5000/api/health\nHeaders: None"
    )

    doc.add_heading("Response:", level=3)
    health_response = doc.add_paragraph(
        "{\n"
        '  "status": "healthy",\n'
        '  "model_loaded": true,\n'
        '  "timestamp": "2025-11-24T10:30:00"\n'
        "}"
    )
    health_response.style = "No Spacing"
    health_response.runs[0].font.name = "Consolas"
    health_response.runs[0].font.size = Pt(9)

    # Endpoint 2: Model Info
    doc.add_heading("9.2 GET /api/model/info", level=2)
    doc.add_paragraph("Returns model metadata and performance metrics.")

    doc.add_heading("Request:", level=3)
    doc.add_paragraph(
        "Method: GET\nURL: http://localhost:5000/api/model/info\nHeaders: None"
    )

    doc.add_heading("Response:", level=3)
    info_response = doc.add_paragraph(
        "{\n"
        '  "model_name": "Random Forest + SMOTE + 20 Features",\n'
        '  "accuracy": 0.7661,\n'
        '  "features": 47,\n'
        '  "training_samples": 5301,\n'
        '  "test_samples": 885,\n'
        '  "classes": ["Dropout", "Enrolled", "Graduate"],\n'
        '  "precision": {"Dropout": 0.82, "Enrolled": 0.47, "Graduate": 0.83},\n'
        '  "recall": {"Dropout": 0.71, "Enrolled": 0.53, "Graduate": 0.86},\n'
        '  "f1_score": {"Dropout": 0.76, "Enrolled": 0.50, "Graduate": 0.84}\n'
        "}"
    )
    info_response.style = "No Spacing"
    info_response.runs[0].font.name = "Consolas"
    info_response.runs[0].font.size = Pt(9)

    # Endpoint 3: Single Prediction
    doc.add_heading("9.3 POST /api/predict", level=2)
    doc.add_paragraph("Predict dropout risk for a single student.")

    doc.add_heading("Request:", level=3)
    doc.add_paragraph(
        "Method: POST\nURL: http://localhost:5000/api/predict\nHeaders: Content-Type: application/json"
    )

    doc.add_heading("Request Body:", level=3)
    single_request = doc.add_paragraph(
        "{\n"
        '  "Marital Status": 1,\n'
        '  "Application Mode": 5,\n'
        '  "Course": 33,\n'
        '  "Daytime/Evening Attendance": 1,\n'
        "  ... (28 required features)\n"
        "}"
    )
    single_request.style = "No Spacing"
    single_request.runs[0].font.name = "Consolas"
    single_request.runs[0].font.size = Pt(9)

    doc.add_heading("Response:", level=3)
    single_response = doc.add_paragraph(
        "{\n"
        '  "prediction": "Graduate",\n'
        '  "confidence": 0.87,\n'
        '  "probabilities": {\n'
        '    "Dropout": 0.08,\n'
        '    "Enrolled": 0.05,\n'
        '    "Graduate": 0.87\n'
        "  },\n"
        '  "risk_level": "Low"\n'
        "}"
    )
    single_response.style = "No Spacing"
    single_response.runs[0].font.name = "Consolas"
    single_response.runs[0].font.size = Pt(9)

    # Endpoint 4: Batch Prediction (Main endpoint)
    doc.add_heading("9.4 POST /api/predict/batch", level=2)
    doc.add_paragraph(
        "Main endpoint for CSV file uploads. Predicts dropout risk for multiple students."
    )

    doc.add_heading("Request:", level=3)
    doc.add_paragraph(
        "Method: POST\nURL: http://localhost:5000/api/predict/batch\nHeaders: Content-Type: multipart/form-data"
    )

    doc.add_heading("Request Body:", level=3)
    doc.add_paragraph("Form Data:\n  file: <CSV file with 28 required columns>")

    doc.add_heading("CSV Format Required Columns:", level=3)
    csv_columns = [
        "Marital Status, Application Mode, Application Order, Course,",
        "Daytime/Evening Attendance, Previous Qualification, Nationality,",
        "Mother's Qualification, Father's Qualification, Mother's Occupation,",
        "Father's Occupation, Displaced, Educational Special Needs, Debtor,",
        "Tuition Fees Up to Date, Gender, Scholarship Holder, Age at Enrollment,",
        "International, Curricular units 1st sem (credited), Curricular units 1st sem (enrolled),",
        "Curricular units 1st sem (evaluations), Curricular units 1st sem (approved),",
        "Curricular units 1st sem (grade), Curricular units 2nd sem (credited),",
        "Curricular units 2nd sem (enrolled), Curricular units 2nd sem (evaluations),",
        "Curricular units 2nd sem (approved), Curricular units 2nd sem (grade)",
    ]
    doc.add_paragraph(" ".join(csv_columns))

    doc.add_heading("Response:", level=3)
    batch_response = doc.add_paragraph(
        "{\n"
        '  "predictions": [\n'
        "    {\n"
        '      "student_id": 1,\n'
        '      "prediction": "Graduate",\n'
        '      "confidence": 0.87,\n'
        '      "probabilities": {"Dropout": 0.08, "Enrolled": 0.05, "Graduate": 0.87},\n'
        '      "risk_level": "Low"\n'
        "    },\n"
        "    { ... more students ... }\n"
        "  ],\n"
        '  "summary": {\n'
        '    "total_students": 5,\n'
        '    "dropout_count": 1,\n'
        '    "enrolled_count": 1,\n'
        '    "graduate_count": 3,\n'
        '    "dropout_percentage": 20.0,\n'
        '    "graduate_percentage": 60.0,\n'
        '    "high_risk_count": 1,\n'
        '    "medium_risk_count": 1,\n'
        '    "low_risk_count": 3\n'
        "  }\n"
        "}"
    )
    batch_response.style = "No Spacing"
    batch_response.runs[0].font.name = "Consolas"
    batch_response.runs[0].font.size = Pt(9)

    # Endpoint 5: Analyze
    doc.add_heading("9.5 POST /api/analyze", level=2)
    doc.add_paragraph(
        "Returns feature importance and analysis for model interpretability."
    )

    doc.add_heading("Request:", level=3)
    doc.add_paragraph(
        "Method: POST\nURL: http://localhost:5000/api/analyze\nHeaders: None"
    )

    doc.add_heading("Response:", level=3)
    analyze_response = doc.add_paragraph(
        "{\n"
        '  "feature_importance": [\n'
        '    {"feature": "Curricular units 2nd sem (approved)", "importance": 0.0686},\n'
        '    {"feature": "completion_rate", "importance": 0.0658},\n'
        "    { ... top 20 features ... }\n"
        "  ],\n"
        '  "model_config": {\n'
        '    "n_estimators": 300,\n'
        '    "max_depth": 20,\n'
        '    "min_samples_split": 10\n'
        "  }\n"
        "}"
    )
    analyze_response.style = "No Spacing"
    analyze_response.runs[0].font.name = "Consolas"
    analyze_response.runs[0].font.size = Pt(9)

    doc.add_heading("Error Responses:", level=2)

    error_codes = [
        ("400 Bad Request", "Missing required fields, invalid CSV format"),
        ("404 Not Found", "Endpoint does not exist"),
        ("500 Internal Server Error", "Model prediction failed, server error"),
    ]

    error_table = doc.add_table(rows=len(error_codes) + 1, cols=2)
    error_table.style = "Light Grid Accent 1"

    header_row = error_table.rows[0]
    header_row.cells[0].text = "Status Code"
    header_row.cells[1].text = "Description"
    for cell in header_row.cells:
        set_cell_background(cell, "C00000")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].bold = True

    for idx, (code, desc) in enumerate(error_codes, 1):
        row = error_table.rows[idx]
        row.cells[0].text = code
        row.cells[1].text = desc

    add_page_break(doc)

    # ==================== 10. DEPLOYMENT GUIDE ====================
    doc.add_heading("10. Deployment Guide", level=1)

    doc.add_paragraph(
        "This section provides step-by-step instructions for deploying MentorAid "
        "in development and production environments."
    )

    doc.add_heading("10.1 Local Development Setup", level=2)

    doc.add_heading("Prerequisites:", level=3)
    prereqs = [
        "Python 3.10 or higher",
        "Node.js 18 or higher",
        "npm 9 or higher",
        "Git for version control",
        "8GB RAM minimum (16GB recommended)",
        "5GB free disk space",
    ]
    for prereq in prereqs:
        doc.add_paragraph(prereq, style="List Bullet")

    doc.add_heading("Backend Setup:", level=3)

    backend_steps = [
        "Navigate to project directory: cd D:\\MentorAid\\MentorAid-main",
        "Run backend startup script: .\\start-backend.ps1",
        "Script automatically:",
        "  - Creates Python virtual environment",
        "  - Installs dependencies from requirements.txt",
        "  - Trains model if not present",
        "  - Starts Flask server on http://localhost:5000",
        "Verify: Open http://localhost:5000/api/health",
        'Expected: {"status": "healthy", "model_loaded": true}',
    ]
    for idx, step in enumerate(backend_steps, 1):
        doc.add_paragraph(
            step, style="List Number" if not step.startswith("  ") else "List Bullet 2"
        )

    doc.add_heading("Frontend Setup:", level=3)

    frontend_steps = [
        "Open new terminal",
        "Navigate to project: cd D:\\MentorAid\\MentorAid-main",
        "Run frontend startup script: .\\start-frontend.ps1",
        "Script automatically:",
        "  - Installs npm dependencies if needed",
        "  - Starts Vite dev server on http://localhost:5173",
        "Verify: Open http://localhost:5173 in browser",
        "Expected: Login page displayed",
    ]
    for idx, step in enumerate(frontend_steps, 1):
        doc.add_paragraph(
            step, style="List Number" if not step.startswith("  ") else "List Bullet 2"
        )

    doc.add_heading("One-Command Deployment:", level=3)
    doc.add_paragraph("For quick setup, run: .\\start-all.ps1")
    doc.add_paragraph(
        "This starts both backend and frontend in separate PowerShell windows."
    )

    doc.add_heading("10.2 Production Deployment", level=2)

    doc.add_heading("Backend Deployment (Cloud Options):", level=3)

    # AWS deployment
    doc.add_paragraph("Option 1: AWS Elastic Beanstalk").bold = True
    aws_steps = [
        "Install AWS CLI and EB CLI",
        "Initialize: eb init -p python-3.10 mentoraid",
        "Create environment: eb create mentoraid-prod",
        "Deploy: eb deploy",
        "Configure environment variables",
        "Set up RDS for database (optional)",
        "Configure CloudWatch for logging",
    ]
    for step in aws_steps:
        doc.add_paragraph(step, style="List Bullet 2")

    # Azure deployment
    doc.add_paragraph("Option 2: Azure App Service").bold = True
    azure_steps = [
        "Install Azure CLI",
        "Login: az login",
        "Create resource group: az group create --name mentoraid-rg --location eastus",
        'Create app service: az webapp create --resource-group mentoraid-rg --plan mentoraid-plan --name mentoraid-api --runtime "PYTHON:3.10"',
        "Deploy: az webapp up --name mentoraid-api",
        "Configure CORS settings",
        "Set up Application Insights",
    ]
    for step in azure_steps:
        doc.add_paragraph(step, style="List Bullet 2")

    # GCP deployment
    doc.add_paragraph("Option 3: Google Cloud Run").bold = True
    gcp_steps = [
        "Install gcloud CLI",
        "Create Dockerfile for Flask app",
        "Build: gcloud builds submit --tag gcr.io/[PROJECT-ID]/mentoraid",
        "Deploy: gcloud run deploy mentoraid --image gcr.io/[PROJECT-ID]/mentoraid --platform managed",
        "Configure environment variables",
        "Set up Cloud Logging",
    ]
    for step in gcp_steps:
        doc.add_paragraph(step, style="List Bullet 2")

    doc.add_heading("Frontend Deployment (Static Hosting):", level=3)

    # Vercel
    doc.add_paragraph("Option 1: Vercel (Recommended)").bold = True
    vercel_steps = [
        "Install Vercel CLI: npm i -g vercel",
        "Build: npm run build",
        "Deploy: vercel --prod",
        "Update API_URL in FileUpload.tsx to production backend URL",
        "Configure environment variables",
        "Set up custom domain (optional)",
    ]
    for step in vercel_steps:
        doc.add_paragraph(step, style="List Bullet 2")

    # Netlify
    doc.add_paragraph("Option 2: Netlify").bold = True
    netlify_steps = [
        "Install Netlify CLI: npm i -g netlify-cli",
        "Build: npm run build",
        "Deploy: netlify deploy --prod --dir=dist",
        "Configure redirects for SPA",
        "Set environment variables",
    ]
    for step in netlify_steps:
        doc.add_paragraph(step, style="List Bullet 2")

    doc.add_heading("10.3 Environment Configuration", level=2)

    doc.add_paragraph("Backend Environment Variables:")
    backend_env = [
        "FLASK_ENV=production",
        "FLASK_DEBUG=0",
        "MODEL_PATH=/path/to/trained_models/",
        "CORS_ORIGINS=https://your-frontend-domain.com",
        "LOG_LEVEL=INFO",
        "MAX_FILE_SIZE=10485760  # 10MB",
    ]
    env_para = doc.add_paragraph()
    env_run = env_para.add_run("\n".join(backend_env))
    env_run.font.name = "Consolas"
    env_run.font.size = Pt(9)

    doc.add_paragraph("\nFrontend Environment Variables:")
    frontend_env = [
        "VITE_API_URL=https://your-backend-domain.com/api",
        "VITE_APP_NAME=MentorAid",
        "VITE_ENV=production",
    ]
    env_para2 = doc.add_paragraph()
    env_run2 = env_para2.add_run("\n".join(frontend_env))
    env_run2.font.name = "Consolas"
    env_run2.font.size = Pt(9)

    doc.add_heading("10.4 Security Considerations", level=2)

    security = [
        "✅ Enable HTTPS for all production endpoints",
        "✅ Implement authentication (JWT tokens recommended)",
        "✅ Add rate limiting to prevent API abuse",
        "✅ Validate all file uploads (size, type, content)",
        "✅ Sanitize user inputs to prevent injection attacks",
        "✅ Set up CORS properly (whitelist specific domains)",
        "✅ Use environment variables for sensitive config",
        "✅ Enable security headers (CSP, X-Frame-Options)",
        "✅ Regular dependency updates for security patches",
        "✅ Monitor logs for suspicious activity",
    ]
    for item in security:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("10.5 Monitoring & Logging", level=2)

    monitoring = [
        "Backend Logging: Use Flask's built-in logging or integrate with CloudWatch/Stackdriver",
        "Frontend Monitoring: Implement error tracking with Sentry or LogRocket",
        "API Monitoring: Track response times, error rates, request volume",
        "Model Monitoring: Log prediction distributions, confidence scores",
        "Alerts: Set up notifications for high error rates or downtime",
        "Performance: Monitor memory usage, CPU, request latency",
    ]
    for item in monitoring:
        doc.add_paragraph(item, style="List Bullet")

    add_page_break(doc)

    # ==================== 11. MODEL STRENGTHS & WEAKNESSES ====================
    doc.add_heading("11. Model Strengths & Weaknesses", level=1)

    doc.add_paragraph(
        "Comprehensive SWOT (Strengths, Weaknesses, Opportunities, Threats) analysis "
        "for each model category to guide model selection and deployment decisions."
    )

    doc.add_heading("11.1 Random Forest + SMOTE (Production Model)", level=2)

    doc.add_heading("Strengths:", level=3)
    rf_strengths = [
        "🏆 Highest accuracy (76.61%) among all 19 models tested",
        "✅ Excellent interpretability via feature importance",
        "✅ Handles non-linear relationships naturally",
        "✅ Robust to overfitting with proper hyperparameters",
        "✅ Fast training (~3 seconds) and prediction (<100ms)",
        "✅ No GPU required, runs on any hardware",
        "✅ SMOTE balancing improves minority class recall (+6% for Dropout)",
        "✅ Feature engineering provides 1.47% accuracy boost",
        "✅ Production-ready with simple deployment (pickle file)",
        "✅ Well-tested and battle-proven algorithm",
    ]
    for strength in rf_strengths:
        doc.add_paragraph(strength, style="List Bullet")

    doc.add_heading("Weaknesses:", level=3)
    rf_weaknesses = [
        "⚠ Enrolled class has lower precision (0.47) - more false positives",
        "⚠ May not capture very complex non-linear patterns",
        "⚠ Model size (~50MB) larger than linear models",
        "⚠ Can be memory-intensive with many trees",
        "⚠ Feature engineering requires domain knowledge",
    ]
    for weakness in rf_weaknesses:
        doc.add_paragraph(weakness, style="List Bullet")

    doc.add_heading("Opportunities:", level=3)
    rf_opportunities = [
        "📈 Collect more data (target 10,000+ students) for better generalization",
        "📈 Implement SHAP values for explainability",
        "📈 A/B test different balancing techniques",
        "📈 Ensemble with other models for marginal gains",
        "📈 Add temporal features (semester-by-semester tracking)",
    ]
    for opp in rf_opportunities:
        doc.add_paragraph(opp, style="List Bullet")

    doc.add_heading("Threats:", level=3)
    rf_threats = [
        "⚠ Data drift over time (academic patterns change)",
        "⚠ Different universities may have different patterns",
        "⚠ Class imbalance may shift in future data",
        "⚠ New features may become available requiring retraining",
    ]
    for threat in rf_threats:
        doc.add_paragraph(threat, style="List Bullet")

    doc.add_heading("11.2 Deep Learning Models", level=2)

    doc.add_heading("Strengths:", level=3)
    dl_strengths = [
        "✅ Can learn complex non-linear patterns",
        "✅ Scalable to larger datasets (20K+ samples)",
        "✅ Automatic feature learning (less manual engineering)",
        "✅ State-of-the-art for image, text, sequential data",
        "✅ Flexible architecture (add layers, change activations)",
    ]
    for strength in dl_strengths:
        doc.add_paragraph(strength, style="List Bullet")

    doc.add_heading("Weaknesses:", level=3)
    dl_weaknesses = [
        "❌ Underperformed Random Forest (74-75% vs 76.61%)",
        "❌ Requires large datasets (20K+ samples minimum)",
        "❌ Longer training time (5-10 minutes vs 3 seconds)",
        "❌ Black box - hard to interpret predictions",
        "❌ Prone to overfitting on small datasets",
        "❌ Requires GPU for efficient training",
        "❌ More hyperparameters to tune",
        "❌ Larger deployment size (TensorFlow runtime)",
    ]
    for weakness in dl_weaknesses:
        doc.add_paragraph(weakness, style="List Bullet")

    doc.add_heading("Recommendation:", level=3)
    doc.add_paragraph(
        "Not recommended for this dataset. Revisit when data grows to 20,000+ samples. "
        "For tabular data with <10K samples, traditional ML (RF, XGBoost) is superior."
    )

    doc.add_heading("11.3 Gradient Boosting Variants (XGBoost, LightGBM)", level=2)

    doc.add_heading("Strengths:", level=3)
    gb_strengths = [
        "✅ Competitive accuracy (74-75%)",
        "✅ Handle missing values automatically",
        "✅ Built-in regularization (L1, L2)",
        "✅ Feature importance available",
        "✅ Fast inference speed",
        "✅ Work well with imbalanced data",
    ]
    for strength in gb_strengths:
        doc.add_paragraph(strength, style="List Bullet")

    doc.add_heading("Weaknesses:", level=3)
    gb_weaknesses = [
        "⚠ Slightly lower accuracy than RF+SMOTE",
        "⚠ More hyperparameters to tune",
        "⚠ Longer training time than Random Forest",
        "⚠ Sequential training (can't parallelize like RF)",
        "⚠ More prone to overfitting without careful tuning",
    ]
    for weakness in gb_weaknesses:
        doc.add_paragraph(weakness, style="List Bullet")

    doc.add_heading("Use Case:", level=3)
    doc.add_paragraph(
        "Good alternative if RF doesn't work. Try if you have missing data or need regularization."
    )

    doc.add_heading(
        "11.4 Simple Models (Logistic Regression, SVM, Naive Bayes)", level=2
    )

    doc.add_heading("Strengths:", level=3)
    simple_strengths = [
        "✅ Extremely fast training (<1 second)",
        "✅ Small model size (<1MB)",
        "✅ Interpretable coefficients (Logistic Regression)",
        "✅ Low computational requirements",
        "✅ Good baseline models",
    ]
    for strength in simple_strengths:
        doc.add_paragraph(strength, style="List Bullet")

    doc.add_heading("Weaknesses:", level=3)
    simple_weaknesses = [
        "❌ Lower accuracy (62-74%)",
        "❌ Can't capture complex interactions",
        "❌ Assume linear decision boundaries (LR, SVM)",
        "❌ Naive Bayes assumes feature independence (violated here)",
        "❌ SVM slow on large datasets",
    ]
    for weakness in simple_weaknesses:
        doc.add_paragraph(weakness, style="List Bullet")

    doc.add_heading("Recommendation:", level=3)
    doc.add_paragraph(
        "Use only as baseline for comparison. Not recommended for production."
    )

    doc.add_heading("11.5 Model Selection Decision Matrix", level=2)

    decision_matrix = [
        ("Scenario", "Recommended Model", "Reason"),
        ("Production (current)", "RF + SMOTE", "Best accuracy, fast, interpretable"),
        ("Limited resources", "Logistic Regression", "Tiny model, fast inference"),
        (
            "Need interpretability",
            "Random Forest",
            "Feature importance, tree visualization",
        ),
        ("Missing data common", "XGBoost/LightGBM", "Handle missing values natively"),
        ("Data grows to 20K+", "Deep Learning", "Scales better with more data"),
        ("Real-time predictions", "Random Forest", "Fast inference (<100ms)"),
        ("Mobile deployment", "Logistic Regression", "Small model size"),
        ("Research/experimentation", "All models", "Compare and validate"),
    ]

    decision_table = doc.add_table(rows=len(decision_matrix), cols=3)
    decision_table.style = "Medium Grid 1 Accent 1"

    for idx, row_data in enumerate(decision_matrix):
        row = decision_table.rows[idx]
        for col_idx, value in enumerate(row_data):
            row.cells[col_idx].text = value
        if idx == 0:
            for cell in row.cells:
                set_cell_background(cell, "2E75B6")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].bold = True

    add_page_break(doc)

    # ==================== 12. RECOMMENDATIONS & FUTURE WORK ====================
    doc.add_heading("12. Recommendations & Future Work", level=1)

    doc.add_paragraph(
        "This section outlines recommendations for production deployment and future enhancements "
        "to improve model performance and system capabilities."
    )

    doc.add_heading("12.1 Immediate Production Recommendations", level=2)

    immediate_recs = [
        "✅ Deploy current RF+SMOTE model (76.61% accuracy is production-ready)",
        "✅ Implement authentication and authorization",
        "✅ Set up database to store predictions and ground truth",
        "✅ Create admin dashboard for monitoring predictions",
        "✅ Implement email notifications for high-risk students",
        "✅ Add user roles (admin, counselor, student)",
        "✅ Set up automated daily prediction runs",
        "✅ Create PDF reports for counselors",
        "✅ Implement audit logging for compliance",
        "✅ Set up automated backups",
    ]
    for rec in immediate_recs:
        doc.add_paragraph(rec, style="List Bullet")

    doc.add_heading("12.2 Model Improvement Roadmap", level=2)

    doc.add_heading("Short-term (1-3 months):", level=3)
    short_term = [
        "Collect ground truth labels from current semester",
        "Retrain model with new data (target 6,000+ samples)",
        "Implement SHAP values for prediction explanations",
        "A/B test different SMOTE variants (ADASYN, BorderlineSMOTE)",
        "Add confidence calibration (Platt scaling)",
        "Create feature engineering pipeline documentation",
        "Implement model versioning and rollback",
        "Set up automated model retraining schedule",
    ]
    for item in short_term:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Medium-term (3-6 months):", level=3)
    medium_term = [
        "Expand dataset to 10,000+ students across multiple semesters",
        "Add temporal features (semester-by-semester tracking)",
        "Implement ensemble of RF + XGBoost for marginal gains",
        "Create student clustering for personalized interventions",
        "Add socioeconomic features if available",
        "Implement time-series analysis for trend detection",
        "Build intervention recommendation system",
        "Create mobile app for counselors",
    ]
    for item in medium_term:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Long-term (6-12 months):", level=3)
    long_term = [
        "Expand to 20,000+ students for deep learning viability",
        "Implement LSTM/Transformer for sequential semester data",
        "Multi-university deployment and transfer learning",
        "Natural language processing on student essays/feedback",
        "Integration with university ERP systems",
        "Automated intervention scheduling",
        "Causal inference analysis (what-if scenarios)",
        "Fairness and bias auditing across demographics",
    ]
    for item in long_term:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("12.3 Data Collection Priorities", level=2)

    data_priorities = [
        ("Priority", "Data Type", "Expected Impact", "Effort"),
        ("High", "Mid-semester grades", "+2-3% accuracy", "Medium"),
        ("High", "Attendance records", "+1-2% accuracy", "Low"),
        ("High", "Previous semester outcomes", "+2-3% accuracy", "Low"),
        ("Medium", "Socioeconomic indicators", "+1% accuracy", "High"),
        ("Medium", "Engagement metrics (LMS)", "+1-2% accuracy", "Medium"),
        ("Medium", "Extracurricular activities", "+0.5-1% accuracy", "Medium"),
        ("Low", "Mental health surveys", "+0.5% accuracy", "Very High"),
        ("Low", "Peer interactions", "+0.5% accuracy", "High"),
    ]

    data_table = doc.add_table(rows=len(data_priorities), cols=4)
    data_table.style = "Light Grid Accent 1"

    for idx, row_data in enumerate(data_priorities):
        row = data_table.rows[idx]
        for col_idx, value in enumerate(row_data):
            row.cells[col_idx].text = value
        if idx == 0:
            for cell in row.cells:
                set_cell_background(cell, "70AD47")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].bold = True

    doc.add_heading("12.4 Technical Debt & Code Quality", level=2)

    tech_debt = [
        "Add comprehensive unit tests (target 80% coverage)",
        "Implement integration tests for API endpoints",
        "Set up CI/CD pipeline (GitHub Actions or Jenkins)",
        "Add type hints to all Python functions",
        "Implement code linting (ESLint for TS, Pylint for Python)",
        "Add API documentation with Swagger/OpenAPI",
        "Implement caching for frequent predictions",
        "Optimize database queries",
        "Add performance profiling",
        "Document all magic numbers and thresholds",
    ]
    for item in tech_debt:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("12.5 Research Opportunities", level=2)

    research = [
        "Fairness Analysis: Ensure model doesn't discriminate by gender, nationality, socioeconomic status",
        "Causal Inference: Use propensity score matching to identify causal factors",
        "Transfer Learning: Test model generalization across different universities",
        "Explainability: Implement LIME/SHAP for individual prediction explanations",
        "Active Learning: Identify most informative students to label next",
        "Multi-task Learning: Jointly predict dropout, GPA, graduation time",
        "Counterfactual Explanations: What changes would prevent dropout?",
        "Uncertainty Quantification: Confidence intervals for predictions",
    ]
    for item in research:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("12.6 Scalability Considerations", level=2)

    scalability = [
        ("Metric", "Current", "Target (1 year)", "Target (3 years)"),
        ("Students", "4,424", "10,000", "50,000+"),
        ("Universities", "1", "3-5", "20+"),
        ("Predictions/day", "<100", "1,000", "10,000+"),
        ("Model retraining", "Manual", "Monthly", "Weekly"),
        ("Response time", "<100ms", "<50ms", "<20ms"),
        ("Uptime", "Dev", "99%", "99.9%"),
    ]

    scale_table = doc.add_table(rows=len(scalability), cols=4)
    scale_table.style = "Light Grid Accent 1"

    for idx, row_data in enumerate(scalability):
        row = scale_table.rows[idx]
        for col_idx, value in enumerate(row_data):
            row.cells[col_idx].text = value
        if idx == 0:
            for cell in row.cells:
                set_cell_background(cell, "FFC000")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].bold = True

    add_page_break(doc)

    # ==================== 13. CONCLUSION ====================
    doc.add_heading("13. Conclusion", level=1)

    doc.add_paragraph(
        "MentorAid represents a comprehensive AI-powered solution for early identification of at-risk "
        "students in higher education. This project successfully demonstrates the application of machine "
        "learning to a critical educational challenge."
    )

    doc.add_heading("13.1 Key Achievements", level=2)

    achievements_final = [
        "🏆 Trained and evaluated 19 machine learning models across 3 development phases",
        "🏆 Achieved 76.61% prediction accuracy with Random Forest + SMOTE + 20 engineered features",
        "🏆 Engineered 20 domain-specific features improving accuracy by 1.47%",
        "🏆 Explored deep learning approaches (TensorFlow/Keras) for comprehensive comparison",
        "🏆 Built production-ready full-stack application (React + TypeScript + Flask)",
        "🏆 Created 5 RESTful API endpoints for seamless integration",
        "🏆 Implemented SMOTE class balancing improving minority class recall by 6-10%",
        "🏆 Achieved fast inference (<100ms per student) suitable for real-time applications",
        "🏆 Documented all 19 models with comprehensive comparisons and SWOT analysis",
        "🏆 Created professional deployment infrastructure with automated startup scripts",
    ]
    for achievement in achievements_final:
        doc.add_paragraph(achievement, style="List Bullet")

    doc.add_heading("13.2 Critical Learnings", level=2)

    learnings = [
        "📚 Feature Engineering Impact: Carefully crafted domain-specific features (+20) provided 1.47% accuracy boost",
        "📚 Simpler is Better: 20 engineered features outperformed 45+ features (noise vs signal)",
        "📚 Traditional ML Excellence: Random Forest (76.61%) beat Deep Learning (74-75%) on small tabular data",
        "📚 Dataset Size Matters: 4,424 samples insufficient for deep learning; need 20K+ for neural networks to excel",
        "📚 Class Balancing Works: SMOTE improved Dropout recall from 65% to 71% (+9.2%)",
        "📚 Overfitting Danger: Aggressive hyperparameter tuning (200 iterations, GridSearch) caused severe overfitting",
        "📚 Ensemble Limitations: Stacking ensemble didn't beat best individual model",
        "📚 Interpretability Value: Feature importance from Random Forest aids understanding and trust",
        "📚 Production Readiness: Fast training (3 sec) and inference (<100ms) enable real-time deployment",
        "📚 Technology Choices: React+TypeScript+Flask+scikit-learn provides optimal balance of power and simplicity",
    ]
    for learning in learnings:
        doc.add_paragraph(learning, style="List Bullet")

    doc.add_heading("13.3 Model Performance Summary", level=2)

    summary_stats = [
        ("Metric", "Value"),
        ("Best Model", "Random Forest + SMOTE + 20 Features"),
        ("Accuracy", "76.61%"),
        ("Dropout Precision/Recall/F1", "0.82 / 0.71 / 0.76"),
        ("Enrolled Precision/Recall/F1", "0.47 / 0.53 / 0.50"),
        ("Graduate Precision/Recall/F1", "0.83 / 0.86 / 0.84"),
        ("Training Time", "~3 seconds"),
        ("Inference Time", "<100ms per student"),
        ("Model Size", "~50MB"),
        ("Features Used", "47 (28 original + 20 engineered - 7 dropped)"),
        ("Training Samples", "5,301 (after SMOTE from 3,539)"),
        ("Test Samples", "885"),
    ]

    summary_table = doc.add_table(rows=len(summary_stats), cols=2)
    summary_table.style = "Medium Grid 1 Accent 1"

    for idx, (metric, value) in enumerate(summary_stats):
        row = summary_table.rows[idx]
        row.cells[0].text = metric
        row.cells[1].text = value
        if idx == 0:
            for cell in row.cells:
                set_cell_background(cell, "2E75B6")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].bold = True
        else:
            set_cell_background(row.cells[0], "E7E6E6")
            row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_heading("13.4 Production Readiness", level=2)

    doc.add_paragraph("MentorAid is production-ready with the following capabilities:")

    prod_ready = [
        "✅ Trained production model with 76.61% accuracy",
        "✅ Complete REST API with 5 endpoints",
        "✅ Modern React frontend with TypeScript type safety",
        "✅ Comprehensive error handling and logging",
        "✅ CORS configured for cross-origin requests",
        "✅ Automated deployment scripts (PowerShell)",
        "✅ API documentation and deployment guides",
        "✅ Sample data and testing infrastructure",
        "✅ Model interpretability via feature importance",
        "✅ Scalable architecture ready for cloud deployment",
    ]
    for item in prod_ready:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("13.5 Impact Potential", level=2)

    doc.add_paragraph(
        "With 76.61% accuracy, MentorAid can correctly identify approximately 680 out of 885 students' "
        "academic outcomes. For the Dropout class specifically:"
    )

    impact_points = [
        "71% recall means catching 71% of at-risk students before they drop out",
        "82% precision means 82% of flagged students are truly at risk (low false positive rate)",
        "Early intervention for 200+ at-risk students per cohort (assuming 283 dropouts per 885 students)",
        "Potential cost savings: $10,000+ per prevented dropout (tuition, recruitment costs)",
        "Improved student outcomes and university retention rates",
        "Data-driven counselor resource allocation",
        "Proactive rather than reactive student support",
    ]
    for point in impact_points:
        doc.add_paragraph(point, style="List Bullet")

    doc.add_heading("13.6 Next Steps", level=2)

    doc.add_paragraph("To maximize MentorAid's impact, we recommend:")

    next_steps_final = [
        "1. Deploy to production with authentication and monitoring",
        "2. Collect ground truth labels from current semester for validation",
        "3. Implement intervention tracking to measure actual impact",
        "4. Expand dataset to 10,000+ students for model improvement",
        "5. Add SHAP explanations for individual student predictions",
        "6. Create counselor training program on using predictions effectively",
        "7. Set up quarterly model retraining with new data",
        "8. Conduct fairness audit across demographic groups",
        "9. Integrate with university ERP systems for automated data feeds",
        "10. Pilot at additional universities for transfer learning validation",
    ]
    for step in next_steps_final:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("13.7 Final Remarks", level=2)

    doc.add_paragraph(
        "This project demonstrates that machine learning can be effectively applied to educational "
        "challenges with modest dataset sizes (4,424 students). The key to success was not complex "
        "deep learning architectures, but rather:"
    )

    success_factors = [
        "Thoughtful feature engineering based on domain knowledge",
        "Appropriate algorithm selection (Random Forest for tabular data)",
        "Class balancing to handle imbalanced outcomes",
        "Systematic model comparison (19 models evaluated)",
        "Focus on interpretability for stakeholder trust",
        "Production-ready software engineering practices",
    ]
    for factor in success_factors:
        doc.add_paragraph(factor, style="List Bullet 2")

    doc.add_paragraph(
        "\n\nMentorAid stands ready to help educational institutions identify at-risk students early, "
        "enabling timely interventions that can change student lives. The combination of strong "
        "predictive performance (76.61%), fast inference (<100ms), interpretable predictions, and "
        "modern web interface positions MentorAid as a practical, deployable solution for improving "
        "student retention and success."
    )

    doc.add_paragraph("\n" + "=" * 70)
    closing = doc.add_paragraph(
        "\nThank you for reviewing this comprehensive documentation."
    )
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "\nFor questions or support, please contact the development team."
    )
    final_para = doc.add_paragraph(
        "\nMentorAid - Empowering Student Success Through AI"
    )
    final_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final_para.runs[0].bold = True
    final_para.runs[0].font.size = Pt(14)
    final_para.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    # Save document
    output_path = os.path.join(
        os.path.dirname(__file__), "MentorAid_Complete_Professional_Documentation.docx"
    )
    doc.save(output_path)
    print(f"✅ Professional documentation created: {output_path}")
    print(f"📄 Total Sections: 13")
    print(f"📊 Total Tables: 20+")
    print(f"📝 Total Pages: ~60-70")
    print(f"🎯 Models Documented: 19")
    print(f"💡 Technology Stack: Fully Explained")
    print(f"🔬 Deep Learning: Included")
    print(f"📈 Comparisons: All vs All")


if __name__ == "__main__":
    create_professional_documentation()
